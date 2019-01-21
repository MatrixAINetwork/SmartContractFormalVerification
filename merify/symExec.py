import base64
import datetime
import pickle
import time
import tokenize
import traceback
import zlib
from collections import namedtuple
from sys import exit
from tokenize import NUMBER, NAME, NEWLINE

from Crypto.Hash import MD5
from Crypto.PublicKey import RSA
from Crypto.Signature import PKCS1_v1_5
from docx import Document
from docx.shared import Inches
from graphviz import Digraph

from analysis import *
from basicblock import BasicBlock
from merify import compare_versions
from test_evm.global_test_params import (UNKNOWN_INSTRUCTION,
                                         EXCEPTION, PICKLE_PATH)
# yzq
from vulnerability import Reentrancy, AssertionFailure, IntegerUnderflow, IntegerOverflow, PrngBug, DosBug, todBug

#logging.basicConfig(level=logging.DEBUG)
log = logging.getLogger()

UNSIGNED_BOUND_NUMBER = 2**256 - 1
CONSTANT_ONES_159 = BitVecVal((1 << 160) - 1, 256)

Assertion = namedtuple('Assertion', ['pc', 'model'])
Underflow = namedtuple('Underflow', ['pc', 'opcode', 'model', 'path'])
Overflow = namedtuple('Overflow', ['pc', 'opcode', 'model', 'path'])
Edge = namedtuple("Edge", ["v1", "v2"]) # Factory Function for tuples is used as dictionary key
SAFE_FLAG = 0
TAINT_FLAG= 1
class Parameter:
    def __init__(self, **kwargs):
        attr_defaults = {
            "stack": [],
            "calls": [],
            "memory": [],
            "visited": [],
            "overflow_pcs": [],
            "underflow_pcs": [],
            "mem": {},
            "analysis": {},
            "sha3_list": {},
            "global_state": {},
            "path_conditions_and_vars": {},
            "taint_path_conditions":[],
            "taint_stack": [],
            "taint_stack_prng": [],
            "taint_memory": [],
            "taint_memory_prng": [],
            "taint_mem": {},
            "taint_mem_prng": {},
            "var_state": {},
            "current_block_path": [],
            "current_path_model":{}
        }
        for (attr, default) in six.iteritems(attr_defaults):
            setattr(self, attr, kwargs.get(attr, default))

    def copy(self):
        _kwargs = custom_deepcopy(self.__dict__)
        return Parameter(**_kwargs)

def initGlobalVars():
    global g_src_map
    global solver
    # Z3 solver
    global separator
    separator = '\\' if sys.platform in ('win32', 'cygwin') else '/'

    if global_params.PARALLEL:
        t2 = Then('simplify', 'solve-eqs', 'smt')
        _t = Then('tseitin-cnf-core', 'split-clause')
        t1 = ParThen(_t, t2)
        solver = OrElse(t1, t2).solver()
    else:
        solver = Solver()

    solver.set("timeout", global_params.TIMEOUT)

    global MSIZE
    MSIZE = False

    global revertible_overflow_pcs
    revertible_overflow_pcs = set()

    global revertible_underflow_pcs
    revertible_underflow_pcs = set()

    global g_disasm_file
    with open(g_disasm_file, 'r') as f:
        disasm = f.read()
    if 'MSIZE' in disasm:
        MSIZE = True

    global g_timeout
    g_timeout = False

    global visited_pcs
    visited_pcs = set()

    global output_number
    output_number = 0

    global results

    global start_block_to_func_sig
    start_block_to_func_sig = {}

    if g_src_map:


        results = {
            'evm_code_coverage': '',
            'vulnerabilities': {
                'integer_underflow': [],
                'integer_overflow': [],
                # 'callstack': [],
                'money_concurrency': [],
                # 'time_dependency': [],
                'reentrancy': [],
                # 'assertion_failure': [],
                # 'parity_multisig_bug_2': [],
                #yzq
                'prng_bug': [],
                'dos_bug': [],
                'tod_bug':[]
            }
        }
    else:
        results = {
            'evm_code_coverage': '',
            'vulnerabilities': {
                'integer_underflow': [],
                'integer_overflow': [],
                'callstack': False,
                'money_concurrency': False,
                'time_dependency': False,
                'reentrancy': False,
                #yzq
                'prng_bug': [],
                'dos_bug':[],
                'tod_bug': []
            }
        }

    global calls_affect_state
    calls_affect_state = {}

    # capturing the last statement of each basic block
    global end_ins_dict
    end_ins_dict = {}

    # capturing all the instructions, keys are corresponding addresses
    global instructions
    instructions = {}

    # capturing the "jump type" of each basic block
    global jump_type
    jump_type = {}

    global vertices
    vertices = {}

    global edges
    edges = {}

    global visited_edges
    visited_edges = {}

    global money_flow_all_paths
    money_flow_all_paths = []

    global reentrancy_all_paths
    reentrancy_all_paths = []

    # store the path condition corresponding to each path in money_flow_all_paths
    global path_conditions
    path_conditions = []
    #yzq
    global global_problematic_pcs
    global_problematic_pcs = {"reentrancy_bug": [], "assertion_failure": [], "integer_underflow": [], "integer_overflow": [], "prng_bug": [], "dos_bug":{}}

    # store global variables, e.g. storage, balance of all paths
    global all_gs
    all_gs = []

    global total_no_of_paths
    total_no_of_paths = 0

    global no_of_test_cases
    no_of_test_cases = 0

    # to generate names for symbolic variables
    global gen
    gen = Generator()


def is_testing_evm():
    return global_params.UNIT_TEST != 0

def compare_storage_and_gas_unit_test(global_state, analysis):
    unit_test = pickle.load(open(PICKLE_PATH, 'rb'))
    test_status = unit_test.compare_with_symExec_result(global_state, analysis)
    exit(test_status)

def change_format():
    with open(g_disasm_file) as disasm_file:
        file_contents = disasm_file.readlines()
        i = 0
        firstLine = file_contents[0].strip('\n')
        for line in file_contents:
            line = line.replace('SELFDESTRUCT', 'SUICIDE')
            line = line.replace('Missing opcode 0xfd', 'REVERT')
            line = line.replace('Missing opcode 0xfe', 'ASSERTFAIL')
            line = line.replace('Missing opcode', 'INVALID')
            line = line.replace(':', '')
            lineParts = line.split(' ')
            try: # removing initial zeroes
                if compare_versions(global_params.EVM_VERSION, global_params.TESTED_EVM_VERSION) > 0:
                    lineParts[0] = str(int(lineParts[0],16))
                else:
                    lineParts[0] = str(int(lineParts[0]))

            except:
                lineParts[0] = lineParts[0]
            lineParts[-1] = lineParts[-1].strip('\n')
            try: # adding arrow if last is a number
                lastInt = lineParts[-1]
                if(int(lastInt, 16) or int(lastInt, 16) == 0) and len(lineParts) > 2:
                    lineParts[-1] = "=>"
                    lineParts.append(lastInt)
            except Exception:
                pass
            file_contents[i] = ' '.join(lineParts)
            i = i + 1
        file_contents[0] = firstLine
        file_contents[-1] += '\n'

    with open(g_disasm_file, 'w') as disasm_file:
        disasm_file.write("\n".join(file_contents))

def build_cfg_and_analyze():
    change_format()
    with open(g_disasm_file, 'r') as disasm_file:
        disasm_file.readline()  # Remove first line
        tokens = tokenize.generate_tokens(disasm_file.readline)
        collect_vertices(tokens)
        construct_bb()
        construct_static_edges()
        full_sym_exec()  # jump targets are constructed on the fly


def print_cfg():
    for block in vertices.values():
        block.display()
    log.info(str(edges))

def print_visited_edges():
    global visited_eduges
    dot = Digraph(comment="visited cfg")
    for edge in visited_edges.keys():
        v1 = edge.v1
        v2 = edge.v2
        if visited_edges[edge] > 0:
            dot.node(str(v2), "visited\n start:"+str(vertices[v2].start)+"\nend:"+str(vertices[v2].end))
        dot.node(str(v1),"visited\n start:"+str(vertices[v1].start)+"\nend:"+str(vertices[v1].end))

    for block in vertices.values():
        for edge in edges[block.start]:
            dot.edge(str(block.start), str(edge))
    dot.render('test-output/visited-cfg.gv', view=False)


def print_dot_cfg():
    global g_src_map
    global instructions
    global jump_type
    dot = Digraph(comment='The Round Table')
    for block in vertices.values():
        if g_src_map:
            if block.start == 0:
                dot.node(str(block.start), "start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n" + "CHECK CALLDATASIZESIZE")
            elif jump_type[block.start] == "terminal":
                dot.node(str(block.start),"start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n"+str(instructions[block.end]))
            else:
                dot.node(str(block.start),"start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n"+g_src_map.get_source_code_for_block(block.start,block.end,instructions))
        else:
            dot.node(str(block.start),str(block.get_instructions()))
        for edge in edges[block.start]:
            dot.edge(str(block.start), str(edge))
    #dot.render('test-output/global-cfg.gv', view=True)
    log.info(dot.source)

#flows:[[1,2,3],[1,2,4]]
def print_moneyflow_dot_path(flows):
    dot = Digraph(comment='The Round Table')
    count = 0
    for path in flows:
        if path:
            dot.node("head" + str(count), "flow:" + str(count))
            dot.edge("head"+str(count),str(path[0]) + str(count))
        for i in range(0,len(path)-1):
            dot.node(str(path[i])+str(count),"pc:"+str(path[i]))
            dot.edge(str(path[i])+str(count), str(path[i+1])+str(count))
        if path:
            dot.node(str(path[-1])+str(count),"pc:"+str(path[-1]))
        count = count + 1

    return dot.render('tmp/bugpath-cfg'+str(output_number)+".gv", view=False, format='png')

#path=[(blockaddress,"content"),(blockaddress:"content")]
#test_case=["variable name":"variable value"]
def print_dot_cfg_path(path, test_case):
    global start_block_to_func_sig
    global output_number
    global g_src_map
    global instructions
    global vertices
    global edges
    output_number=output_number+1
    #path[-1] = (path[-1][0], path[-1][1] + "\n branch test case:\n" + str(test_case))
    dot = Digraph(comment='BUG PATH')
    #dot.node("test case", str(test_case))
    for block in vertices.values():
        for edge in edges[block.start]:
            dot.edge(str(block.start), str(edge))
    for block_pair in path:
        if g_src_map:
            block = vertices[block_pair[0]]
            if block.start == 0:
                dot.node(str(block.start), "start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n" + "CHECK CALLDATASIZESIZE")
            elif jump_type[block.start] == "terminal":
                dot.node(str(block.start),"start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n"+str(instructions[block.end]))
            else:
                dot.node(str(block.start),"start:"+str(block.start)+"\n end:"+str(block.end) +\
                     "\n source:\n"+g_src_map.get_source_code_for_block(block.start,block.end,instructions))
        else:
            if block_pair[0] in start_block_to_func_sig.keys():
                dot.node(str(block_pair[0]), "block number:" + str(block_pair[0]) + ":\n" + str(block_pair[1]) + "\n functioin hash:" + str(start_block_to_func_sig[block_pair[0]]))
            else:
                dot.node(str(block_pair[0]), "block number:" + str(block_pair[0])+":\n" + str(block_pair[1]))
    return dot.render('tmp/bugpath-cfg'+str(output_number)+".gv", view=False, format='png')
    #log.info(dot.source)

def mapping_push_instruction(current_line_content, current_ins_address, idx, positions, length):
    global g_src_map

    while (idx < length):
        if not positions[idx]:
            return idx + 1
        name = positions[idx]['name']
        if name.startswith("tag"):
            idx += 1
        else:
            if name.startswith("PUSH"):
                if name == "PUSH":
                    value = positions[idx]['value']
                    instr_value = current_line_content.split(" ")[1]
                    if int(value, 16) == int(instr_value, 16):
                        g_src_map.instr_positions[current_ins_address] = g_src_map.positions[idx]
                        idx += 1
                        break;
                    else:
                        raise Exception("Source map error")
                else:
                    g_src_map.instr_positions[current_ins_address] = g_src_map.positions[idx]
                    idx += 1
                    break;
            else:
                raise Exception("Source map error")
    return idx

def mapping_non_push_instruction(current_line_content, current_ins_address, idx, positions, length):
    global g_src_map

    while (idx < length):
        if not positions[idx]:
            return idx + 1
        name = positions[idx]['name']
        if name.startswith("tag"):
            idx += 1
        else:
            instr_name = current_line_content.split(" ")[0]
            if name == instr_name or name == "INVALID" and instr_name == "ASSERTFAIL" or name == "KECCAK256" and instr_name == "SHA3" or name == "SELFDESTRUCT" and instr_name == "SUICIDE":
                g_src_map.instr_positions[current_ins_address] = g_src_map.positions[idx]
                idx += 1
                break;
            else:
                raise Exception("Source map error")
    return idx

# 1. Parse the disassembled file
# 2. Then identify each basic block (i.e. one-in, one-out)
# 3. Store them in vertices
def collect_vertices(tokens):
    global g_src_map
    if g_src_map:
        idx = 0
        positions = g_src_map.positions
        length = len(positions)
    global end_ins_dict
    global instructions
    global jump_type

    current_ins_address = 0
    last_ins_address = 0
    is_new_line = True
    current_block = 0
    current_line_content = ""
    wait_for_push = False
    is_new_block = False
    #log.info("<<<<<<<<<")
    for tok_type, tok_string, (srow, scol), mi, line_number in tokens:
        # printtoken(tok_type,tok_string,(srow,scol),mi,line_number)
        #log.info(str(line_number))
        if wait_for_push is True:
            push_val = ""
            for ptok_type, ptok_string, _, _, _ in tokens:
                if ptok_type == NEWLINE:
                    is_new_line = True
                    current_line_content += push_val + ' '
                    instructions[current_ins_address] = current_line_content
                    idx = mapping_push_instruction(current_line_content, current_ins_address, idx, positions, length) if g_src_map else None
                    log.debug(current_line_content)
                    current_line_content = ""
                    wait_for_push = False
                    break
                try:
                    int(ptok_string, 16)
                    push_val += ptok_string
                except ValueError:
                    pass

            continue
        elif is_new_line is True and tok_type == NUMBER:  # looking for a line number
            last_ins_address = current_ins_address
            try:
                current_ins_address = int(tok_string)
            except ValueError:
                log.critical("ERROR when parsing row %d col %d", srow, scol)
                quit()
            is_new_line = False
            if is_new_block:
                current_block = current_ins_address
                is_new_block = False
            continue
        elif tok_type == NEWLINE:
            is_new_line = True
            log.debug(current_line_content)
            instructions[current_ins_address] = current_line_content
            idx = mapping_non_push_instruction(current_line_content, current_ins_address, idx, positions, length) if g_src_map else None
            current_line_content = ""
            continue
        elif tok_type == NAME:
            if tok_string == "JUMPDEST":
                if last_ins_address not in end_ins_dict:
                    end_ins_dict[current_block] = last_ins_address
                current_block = current_ins_address
                is_new_block = False
            elif tok_string == "STOP" or tok_string == "RETURN" or tok_string == "SUICIDE" or tok_string == "REVERT" or tok_string == "ASSERTFAIL":
                jump_type[current_block] = "terminal"
                end_ins_dict[current_block] = current_ins_address
            elif tok_string == "JUMP":
                jump_type[current_block] = "unconditional"
                end_ins_dict[current_block] = current_ins_address
                is_new_block = True
            elif tok_string == "JUMPI":
                jump_type[current_block] = "conditional"
                end_ins_dict[current_block] = current_ins_address
                is_new_block = True
            elif tok_string.startswith('PUSH', 0):
                wait_for_push = True
            is_new_line = False
        if tok_string != "=" and tok_string != ">":
            current_line_content += tok_string + " "
    #log.info("<<<<<<<<<<<<<<<<")
    #log.info("**************")
    # for i in instructions.keys():
    #     log.info(str(i) + ":" + instructions[i])
    # log.info("*************")
    if current_block not in end_ins_dict:
        log.debug("current block: %d", current_block)
        log.debug("last line: %d", current_ins_address)
        end_ins_dict[current_block] = current_ins_address

    if current_block not in jump_type:
        jump_type[current_block] = "terminal"

    for key in end_ins_dict:
        if key not in jump_type:
            jump_type[key] = "falls_to"


def construct_bb():
    global vertices
    global edges
    sorted_addresses = sorted(instructions.keys())
    size = len(sorted_addresses)
    for key in end_ins_dict:
        end_address = end_ins_dict[key]
        block = BasicBlock(key, end_address)
        if key not in instructions:
            continue
        block.add_instruction(instructions[key])
        i = sorted_addresses.index(key) + 1
        while i < size and sorted_addresses[i] <= end_address:
            block.add_instruction(instructions[sorted_addresses[i]])
            i += 1
        block.set_block_type(jump_type[key])
        vertices[key] = block
        edges[key] = []


def construct_static_edges():
    add_falls_to()  # these edges are static


def add_falls_to():
    global vertices
    global edges
    key_list = sorted(jump_type.keys())
    length = len(key_list)
    for i, key in enumerate(key_list):
        if jump_type[key] != "terminal" and jump_type[key] != "unconditional" and i+1 < length:
            target = key_list[i+1]
            edges[key].append(target)
            vertices[target].set_jump_from(key)
            vertices[key].set_falls_to(target)


def get_init_global_state(path_conditions_and_vars):
    global_state = {"balance" : {}, "pc": 0}
    init_is = init_ia = deposited_value = sender_address = receiver_address = gas_price = origin = currentCoinbase = currentNumber = currentDifficulty = currentGasLimit = callData = None

    if global_params.INPUT_STATE:
        with open('state.json') as f:
            state = json.loads(f.read())
            if state["Is"]["balance"]:
                init_is = int(state["Is"]["balance"], 16)
            if state["Ia"]["balance"]:
                init_ia = int(state["Ia"]["balance"], 16)
            if state["exec"]["value"]:
                deposited_value = 0
            if state["Is"]["address"]:
                sender_address = int(state["Is"]["address"], 16)
            if state["Ia"]["address"]:
                receiver_address = int(state["Ia"]["address"], 16)
            if state["exec"]["gasPrice"]:
                gas_price = int(state["exec"]["gasPrice"], 16)
            if state["exec"]["origin"]:
                origin = int(state["exec"]["origin"], 16)
            if state["env"]["currentCoinbase"]:
                currentCoinbase = int(state["env"]["currentCoinbase"], 16)
            if state["env"]["currentNumber"]:
                currentNumber = int(state["env"]["currentNumber"], 16)
            if state["env"]["currentDifficulty"]:
                currentDifficulty = int(state["env"]["currentDifficulty"], 16)
            if state["env"]["currentGasLimit"]:
                currentGasLimit = int(state["env"]["currentGasLimit"], 16)

    # for some weird reason these 3 vars are stored in path_conditions insteaad of global_state
    else:
        sender_address = BitVec("Is", 256)
        receiver_address = BitVec("Ia", 256)
        deposited_value = BitVec("Iv", 256)
        init_is = BitVec("init_Is", 256)
        init_ia = BitVec("init_Ia", 256)

    path_conditions_and_vars["Is"] = sender_address
    path_conditions_and_vars["Ia"] = receiver_address
    path_conditions_and_vars["Iv"] = deposited_value

    constraint = (deposited_value >= BitVecVal(0, 256))
    path_conditions_and_vars["path_condition"].append(constraint)
    constraint = (init_is >= deposited_value)
    path_conditions_and_vars["path_condition"].append(constraint)
    constraint = (init_ia >= BitVecVal(0, 256))
    path_conditions_and_vars["path_condition"].append(constraint)

    # update the balances of the "caller" and "callee"

    global_state["balance"]["Is"] = (init_is - deposited_value)
    global_state["balance"]["Ia"] = (init_ia + deposited_value)

    if not gas_price:
        new_var_name = gen.gen_gas_price_var()
        gas_price = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = gas_price

    if not origin:
        new_var_name = gen.gen_origin_var()
        origin = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = origin

    if not currentCoinbase:
        new_var_name = "IH_c"
        currentCoinbase = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = currentCoinbase

    if not currentNumber:
        new_var_name = "IH_i"
        currentNumber = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = currentNumber

    if not currentDifficulty:
        new_var_name = "IH_d"
        currentDifficulty = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = currentDifficulty

    if not currentGasLimit:
        new_var_name = "IH_l"
        currentGasLimit = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = currentGasLimit

    new_var_name = "IH_s"
    currentTimestamp = BitVec(new_var_name, 256)
    path_conditions_and_vars[new_var_name] = currentTimestamp

    # the state of the current current contract
    if "Ia" not in global_state:
        global_state["Ia"] = {}
    global_state["miu_i"] = 0
    global_state["value"] = deposited_value
    global_state["sender_address"] = sender_address
    global_state["receiver_address"] = receiver_address
    global_state["gas_price"] = gas_price
    global_state["origin"] = origin
    global_state["currentCoinbase"] = currentCoinbase
    global_state["currentTimestamp"] = currentTimestamp
    global_state["currentNumber"] = currentNumber
    global_state["currentDifficulty"] = currentDifficulty
    global_state["currentGasLimit"] = currentGasLimit
    global_state["lastStore"] = {}
    global_state["flowSload"] = {}

    return global_state

def get_start_block_to_func_sig():
    state = 0
    func_sig = None
    for pc in sorted(instructions.keys()):
        instr = instructions[pc]
        if state == 0 and instr.startswith('PUSH4'):
            state += 1
            func_sig = instr.split(' ')[1][2:]
        elif state == 1 and instr.startswith('EQ'):
            state += 1
        elif state == 2 and instr.startswith('PUSH'):
            state = 0
            pc = instr.split(' ')[1]
            pc = int(pc, 16)
            start_block_to_func_sig[pc] = func_sig
        else:
            state = 0
    return start_block_to_func_sig

def full_sym_exec():
    # executing, starting from beginning
    path_conditions_and_vars = {"path_condition" : []}
    global_state = get_init_global_state(path_conditions_and_vars)
    analysis = init_analysis()
    params = Parameter(path_conditions_and_vars=path_conditions_and_vars, global_state=global_state, analysis=analysis)
    #if g_src_map:
    get_start_block_to_func_sig()
    return sym_exec_block(params, 0, 0, 0, -1, 'fallback')


# Symbolically executing a block from the start address
def sym_exec_block(params, block, pre_block, depth, func_call, current_func_name):
    global solver
    global visited_edges
    global money_flow_all_paths
    global path_conditions
    global global_problematic_pcs
    global all_gs
    global results
    global g_src_map


    visited = params.visited
    stack = params.stack
    mem = params.mem
    memory = params.memory
    global_state = params.global_state
    sha3_list = params.sha3_list
    path_conditions_and_vars = params.path_conditions_and_vars
    analysis = params.analysis
    calls = params.calls
    overflow_pcs = params.overflow_pcs
    underflow_pcs = params.underflow_pcs
    taint_stack = params.taint_stack
    taint_stack_prng = params.taint_stack_prng
    taint_mem = params.taint_mem
    taint_mem_prng = params.taint_mem_prng
    taint_memory = params.taint_memory
    taint_memory_prng = params.taint_memory_prng
    var_state = params.var_state


    if block < 0:
        log.debug("UNKNOWN JUMP ADDRESS. TERMINATING THIS PATH")
        return ["ERROR"]

    log.debug("Reach block address %d \n", block)

    if g_src_map:


        if block in start_block_to_func_sig:
            func_sig = start_block_to_func_sig[block]
            current_func_name = g_src_map.sig_to_func[func_sig]
            pattern = r'(\w[\w\d_]*)\((.*)\)$'
            match = re.match(pattern, current_func_name)
            if match:
                current_func_name =  list(match.groups())[0]

    current_edge = Edge(pre_block, block)
    if current_edge in visited_edges:
        updated_count_number = visited_edges[current_edge] + 1
        visited_edges.update({current_edge: updated_count_number})
    else:
        visited_edges.update({current_edge: 1})

    if visited_edges[current_edge] > global_params.LOOP_LIMIT:
        if check_loop_limit(params.current_block_path):
            log.debug("Overcome a number of loop limit. Terminating this path ...")
            if jump_type[pre_block] == "conditional":
                for iblock in params.current_block_path:
                    if iblock[0] in start_block_to_func_sig.keys():
                        dos_pc = iblock[0]
                        if not dos_pc in global_problematic_pcs["dos_bug"]:
                            global_problematic_pcs["dos_bug"][dos_pc] = dos_pc
                            global_params.BUG_TEST_CASE['DOS'][dos_pc] = params.current_path_model
                            global_params.BUG_PC['DOS'][dos_pc] = str(dos_pc) + ": " + str(
                                instructions[dos_pc]) + " :dos caused by loop gas limit"
                            global_params.BUG_BLOCK_PATH['DOS'][dos_pc] = params.current_block_path
            return stack

    current_gas_used = analysis["gas"]
    if current_gas_used > global_params.GAS_LIMIT:
        log.debug("Run out of gas. Terminating this path ... ")
        return stack

    # Execute every instruction, one at a time
    try:
        block_ins = vertices[block].get_instructions()
    except KeyError:
        #print(str(KeyError))
        log.debug("This path results in an exception, possibly an invalid jump address")
        return ["ERROR"]
    for instr in block_ins:
        sym_exec_ins(params, block, instr, func_call, current_func_name)
        instr_parts = str.split(instr, ' ')
        opcode = instr_parts[0]
        if (opcode!= "CALLDATASIZE") and (opcode != "CALLDATALOAD") and (global_params.PRE_OPCODE_COUNT >0 ):
            global_params.PRE_OPCODE_COUNT = global_params.PRE_OPCODE_COUNT - 1

    # Mark that this basic block in the visited blocks
    visited.append(block)
    depth += 1

    reentrancy_all_paths.append(analysis["reentrancy_bug"])


    # Go to next Basic Block(s)
    if jump_type[block] == "terminal" or depth > global_params.DEPTH_LIMIT:
        global total_no_of_paths
        global no_of_test_cases
        # add tod
        for store_address in global_state["lastStore"]:
            if not store_address in global_params.GLOBAL_PC_SSTORE:
                global_params.GLOBAL_PC_SSTORE[store_address] = {}
            sstore_pc = global_state["lastStore"][store_address]["pc_num"]
            if not sstore_pc in global_params.GLOBAL_PC_SSTORE[store_address]:
                global_params.GLOBAL_PC_SSTORE[store_address][sstore_pc] = []
            # for condition in global_state["lastStore"][store_address]["path_condition"]:
            global_params.GLOBAL_PC_SSTORE[store_address][sstore_pc] = global_state["lastStore"][store_address]["path_condition"]
        global_state["lastStore"].clear()
        global_state["flowSload"].clear()
        total_no_of_paths += 1

        if global_params.GENERATE_TEST_CASES:
            try:
                model = solver.model()
                no_of_test_cases += 1
                filename = "test%s.otest" % no_of_test_cases
                with open(filename, 'w') as f:
                    for variable in model.decls():
                        f.write(str(variable) + " = " + str(model[variable]) + "\n")
                if os.stat(filename).st_size == 0:
                    os.remove(filename)
                    no_of_test_cases -= 1
            except Exception as e:
                #print(str(Exception))
                Exception('Unknown Jump-Type')
                pass

        log.debug("TERMINATING A PATH ...")
        if is_testing_evm():
            compare_storage_and_gas_unit_test(global_state, analysis)

    elif jump_type[block] == "unconditional":  # executing "JUMP"
        successor = vertices[block].get_jump_target()
        new_params = params.copy()
        new_params.global_state["pc"] = successor
        if g_src_map:
            source_code = g_src_map.get_source_code(global_state['pc'])
            if source_code in g_src_map.func_call_names:
                func_call = global_state['pc']
        new_params.current_block_path.append((successor,"unconditionnal branch expression"))
        sym_exec_block(new_params, successor, block, depth, func_call, current_func_name)
    elif jump_type[block] == "falls_to":  # just follow to the next basic block
        successor = vertices[block].get_falls_to()
        new_params = params.copy()
        new_params.global_state["pc"] = successor
        new_params.current_block_path.append((successor,"falls_to"))
        sym_exec_block(new_params, successor, block, depth, func_call, current_func_name)
    elif jump_type[block] == "conditional":  # executing "JUMPI"

        # A choice point, we proceed with depth first search

        branch_expression = vertices[block].get_branch_expression()

        log.debug("Branch expression: " + str(branch_expression))

        solver.push()  # SET A BOUNDARY FOR SOLVER
        solver.add(branch_expression)

        try:
            if solver.check() == unsat:
                log.debug("INFEASIBLE PATH DETECTED")
            else:
                left_branch = vertices[block].get_jump_target()
                new_params = params.copy()
                model = solver.model()
                for variable in model.decls():
                    new_params.current_path_model[str(variable)]=str(model[variable])
                new_params.global_state["pc"] = left_branch
                new_params.path_conditions_and_vars["path_condition"].append(branch_expression)
                new_params.taint_path_conditions.insert(0,vertices[block].get_taint_branch_expression())
                last_idx = len(new_params.path_conditions_and_vars["path_condition"]) - 1
                new_params.current_block_path.append((left_branch,("conditionnal True branch: \n" + "branch expression:" + str(branch_expression))))
                sym_exec_block(new_params, left_branch, block, depth, func_call, current_func_name)
        except TimeoutError:
            raise
        except Exception as e:
            #print(str(Exception))
            if global_params.DEBUG_MODE:
                traceback.print_exc()

        solver.pop()  # POP SOLVER CONTEXT

        solver.push()  # SET A BOUNDARY FOR SOLVER
        negated_branch_expression = Not(branch_expression)
        solver.add(negated_branch_expression)

        log.debug("Negated branch expression: " + str(negated_branch_expression))

        try:
            if solver.check() == unsat:
                # Note that this check can be optimized. I.e. if the previous check succeeds,
                # no need to check for the negated condition, but we can immediately go into
                # the else branch
                log.debug("INFEASIBLE PATH DETECTED")
            else:
                right_branch = vertices[block].get_falls_to()
                new_params = params.copy()
                model = solver.model()
                for variable in model.decls():
                    new_params.current_path_model[str(variable)]=str(model[variable])
                new_params.global_state["pc"] = right_branch
                new_params.path_conditions_and_vars["path_condition"].append(negated_branch_expression)
                new_params.taint_path_conditions.insert(0, vertices[block].get_taint_branch_expression())
                last_idx = len(new_params.path_conditions_and_vars["path_condition"]) - 1
                new_params.current_block_path.append((right_branch, ("conditionnal False branch: \n" + "branch expression:" + str(negated_branch_expression))))
                sym_exec_block(new_params, right_branch, block, depth, func_call, current_func_name)
        except TimeoutError:
            raise
        except Exception as e:
            #print(str(Exception))
            if global_params.DEBUG_MODE:
                traceback.print_exc()
        solver.pop()  # POP SOLVER CONTEXT
        updated_count_number = visited_edges[current_edge] - 1
        visited_edges.update({current_edge: updated_count_number})
    else:
        updated_count_number = visited_edges[current_edge] - 1
        visited_edges.update({current_edge: updated_count_number})
        #print(str(Exception('Unknown Jump-Type')))
        raise Exception('Unknown Jump-Type')

def check_loop_limit(current_path):
    current_loop = 0
    length = len(current_path)

    if length >  2:
        snewblock = current_path[-3][0]
        soldblock = current_path[-2][0]
        newblock = current_path[-2][0]
        oldblock = current_path[-1][0]
        for i in range(3,length):
            oldblock = newblock
            newblock = current_path[-i][0]
            if oldblock == soldblock and newblock == snewblock:
                current_loop += 1
        if current_loop >= global_params.LOOP_LIMIT:
            return True
    else:
        return False




# Symbolically executing an instruction
def sym_exec_ins(params, block, instr, func_call, current_func_name):
    global MSIZE
    global visited_pcs
    global solver
    global vertices
    global edges
    global g_src_map
    global calls_affect_state

    stack = params.stack
    mem = params.mem
    memory = params.memory
    global_state = params.global_state
    sha3_list = params.sha3_list
    path_conditions_and_vars = params.path_conditions_and_vars
    taint_path_conditions=params.taint_path_conditions
    analysis = params.analysis
    calls = params.calls
    overflow_pcs = params.overflow_pcs
    underflow_pcs = params.underflow_pcs
    taint_memory = params.taint_memory
    taint_memory_prng = params.taint_memory_prng
    taint_stack = params.taint_stack
    taint_stack_prng = params.taint_stack_prng
    taint_mem = params.taint_mem
    taint_mem_prng = params.taint_mem_prng
    current_block_path = params.current_block_path
    current_path_model = params.current_path_model
    visited_pcs.add(global_state["pc"])

    instr_parts = str.split(instr, ' ')
    opcode = instr_parts[0]
    #if opcode == "CALL":
        #print(str(global_state["pc"])+": "+opcode)
    if opcode == "INVALID":
        return
    elif opcode == "ASSERTFAIL":
        if check_sat(solver, False) != unsat:
            model = solver.model()
        if g_src_map:
            source_code = g_src_map.get_source_code(global_state['pc'])
            source_code = source_code.split("(")[0]
            func_name = source_code.strip()
            if func_name == "assert":
                global_problematic_pcs["assertion_failure"].append(global_state["pc"])
                global_params.BUG_PC['ASSERTFAIL'][global_state['pc']] = str(global_state['pc']) + ": " + str(
                    opcode) + ": invalid assertfail"
                global_params.BUG_TEST_CASE['ASSERTFAIL'][global_state['pc']] = current_path_model
                global_params.BUG_BLOCK_PATH['ASSERTFAIL'][global_state['pc']] = current_block_path
            elif func_call != -1:
                global_problematic_pcs["assertion_failure"].append( model)
                global_params.BUG_PC['ASSERTFAIL'][func_call] = str(func_call) + ": " + str(
                    opcode) + ": invalid assertfail"
                global_params.BUG_TEST_CASE['ASSERTFAIL'][func_call] = current_path_model
                global_params.BUG_BLOCK_PATH['ASSERTFAIL'][func_call] = current_block_path
        else:
            global_problematic_pcs["assertion_failure"].append(global_state["pc"])
            global_params.BUG_PC['ASSERTFAIL'][global_state['pc']]=str(global_state['pc']) + ": "+ str(opcode) + ": invalid assertfail"
            global_params.BUG_TEST_CASE['ASSERTFAIL'][global_state['pc']]=current_path_model
            global_params.BUG_BLOCK_PATH['ASSERTFAIL'][global_state['pc']]=current_block_path
        return

    # collecting the analysis result by calling this skeletal function
    # this should be done before symbolically executing the instruction,
    # since SE will modify the stack and mem
    update_analysis(analysis, opcode, stack, mem, global_state, path_conditions_and_vars, taint_path_conditions,solver,taint_stack,taint_stack_prng,current_block_path,current_path_model)
    if opcode == "CALL" and analysis["reentrancy_bug"] and analysis["reentrancy_bug"][-1]:
        global_problematic_pcs["reentrancy_bug"].append(global_state["pc"])
    #yzq:
    if (opcode == "CALL" or opcode == "SSTORE") and analysis["prng_bug"] and analysis["prng_bug"][-1]:
        global_problematic_pcs["prng_bug"].append(global_state["pc"])
    if opcode == "JUMPI"  and analysis["dos_bug"][-1]:
        global_problematic_pcs["dos_bug"][global_state["pc"]] = block
    #print("***analysis:"+str(analysis))
    log.debug("==============================")
    log.debug("EXECUTING: " + instr)
    #print("opcode:"+str(opcode))
    #print("pc:"+str(global_state["pc"]))
    #print("ts:"+str(taint_stack_prng))
    #
    #  0s: Stop and Arithmetic Operations
    #
    if opcode == "STOP":
        global_state["pc"] = global_state["pc"] + 1
        return
    elif opcode == "ADD":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            # Type conversion is needed when they are mismatched
            if isReal(first) and isSymbolic(second):
                first = BitVecVal(first, 256)
                computed = first + second
            elif isSymbolic(first) and isReal(second):
                second = BitVecVal(second, 256)
                computed = first + second
            else:
                # both are real and we need to manually modulus with 2 ** 256
                # if both are symbolic z3 takes care of modulus automatically
                computed = (first + second) % (2 ** 256)
            computed = simplify(computed) if is_expr(computed) else computed

            check_revert = False
            if jump_type[block] == 'conditional':
                jump_target = vertices[block].get_jump_target()
                falls_to = vertices[block].get_falls_to()
                check_revert = any([True for instruction in vertices[jump_target].get_instructions() if
                                    instruction.startswith('REVERT')])
                if not check_revert:
                    check_revert = any([True for instruction in vertices[falls_to].get_instructions() if
                                        instruction.startswith('REVERT')])

            if (global_params.PRE_OPCODE_COUNT <= 0) and (jump_type[block] != 'conditional' or not check_revert) and (
                    taint_first | taint_second):
                if not isAllReal(computed, first):
                    solver.push()
                    solver.add(UGT(first, computed))
                    if check_sat(solver) == sat:
                        global_problematic_pcs['integer_overflow'].append(Overflow(global_state['pc'] - 1, opcode, solver.model(), current_block_path))
                        overflow_pcs.append(global_state['pc'] - 1)
                    solver.pop()

            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MUL":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isReal(first) and isSymbolic(second):
                first = BitVecVal(first, 256)
            elif isSymbolic(first) and isReal(second):
                second = BitVecVal(second, 256)
            computed = first * second & UNSIGNED_BOUND_NUMBER
            computed = simplify(computed) if is_expr(computed) else computed

            check_revert = False
            if jump_type[block] == 'conditional':
                jump_target = vertices[block].get_jump_target()
                falls_to = vertices[block].get_falls_to()
                check_revert = any([True for instruction in vertices[jump_target].get_instructions() if
                                    instruction.startswith('REVERT')])
                if not check_revert:
                    check_revert = any([True for instruction in vertices[falls_to].get_instructions() if
                                        instruction.startswith('REVERT')])

            if (jump_type[block] != 'conditional') and (taint_first | taint_second):
                if taint_first or taint_second:
                    if not isAllReal(computed, first):
                        if isReal(first):
                            solver.push()
                            solver.add(Not(computed / first == second))
                            if check_sat(solver) == sat:
                                global_problematic_pcs['integer_overflow'].append(Overflow(global_state['pc'] - 1, opcode, solver.model(), current_block_path))
                                overflow_pcs.append(global_state['pc'] - 1)
                            solver.pop()
                        if isReal(second):
                            solver.push()
                            solver.add(Not(computed / second == first))
                            if check_sat(solver) == sat:
                                global_problematic_pcs['integer_overflow'].append(Overflow(global_state['pc'] - 1, opcode, solver.model(), current_block_path))
                                overflow_pcs.append(global_state['pc'] - 1)
                            solver.pop()
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SUB":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isReal(first) and isSymbolic(second):
                first = BitVecVal(first, 256)
                computed = first - second
            elif isSymbolic(first) and isReal(second):
                second = BitVecVal(second, 256)
                computed = first - second
            else:
                computed = (first - second) % (2 ** 256)
            computed = simplify(computed) if is_expr(computed) else computed

            check_revert = False
            if jump_type[block] == 'conditional':
                jump_target = vertices[block].get_jump_target()
                falls_to = vertices[block].get_falls_to()
                check_revert = any([True for instruction in vertices[jump_target].get_instructions() if
                                    instruction.startswith('REVERT')])
                if not check_revert:
                    check_revert = any([True for instruction in vertices[falls_to].get_instructions() if
                                        instruction.startswith('REVERT')])

            if (global_params.PRE_OPCODE_COUNT <= 0) and (jump_type[block] != 'conditional' or not check_revert) and (
                    taint_first | taint_second):
                if not isAllReal(first, second):
                    solver.push()
                    solver.add(UGT(second, first))
                    if check_sat(solver) == sat:
                        global_problematic_pcs['integer_underflow'].append(Underflow(global_state['pc'] - 1, opcode, solver.model(), current_block_path))
                        underflow_pcs.append(global_state['pc'] - 1)
                    solver.pop()

            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "DIV":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if second == 0:
                    computed = 0
                else:
                    first = to_unsigned(first)
                    second = to_unsigned(second)
                    computed = first / second
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add( Not (second == 0) )
                if check_sat(solver) == unsat:
                    computed = 0 #div zero with no exception but make the computed result to 0
                else:
                    computed = UDiv(first, second)
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0, taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SDIV":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                first = to_signed(first)
                second = to_signed(second)
                if second == 0:
                    computed = 0
                elif first == -2**255 and second == -1:
                    computed = -2**255
                else:
                    sign = -1 if (first / second) < 0 else 1
                    computed = sign * ( abs(first) / abs(second) )
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add(Not(second == 0))
                if check_sat(solver) == unsat:
                    computed = 0
                else:
                    solver.push()
                    solver.add( Not( And(first == -2**255, second == -1 ) ))
                    if check_sat(solver) == unsat:
                        computed = -2**255
                    else:
                        solver.push()
                        solver.add(first / second < 0)
                        sign = -1 if check_sat(solver) == sat else 1
                        z3_abs = lambda x: If(x >= 0, x, -x)
                        first = z3_abs(first)
                        second = z3_abs(second)
                        computed = sign * (first / second)
                        solver.pop()
                    solver.pop()
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MOD":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if second == 0:
                    computed = 0
                else:
                    first = to_unsigned(first)
                    second = to_unsigned(second)
                    computed = first % second & UNSIGNED_BOUND_NUMBER

            else:
                first = to_symbolic(first)
                second = to_symbolic(second)

                solver.push()
                solver.add(Not(second == 0))
                if check_sat(solver) == unsat:
                    # it is provable that second is indeed equal to zero
                    computed = 0
                else:
                    computed = URem(first, second)
                solver.pop()

            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SMOD":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if second == 0:
                    computed = 0
                else:
                    first = to_signed(first)
                    second = to_signed(second)
                    sign = -1 if first < 0 else 1
                    computed = sign * (abs(first) % abs(second))
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)

                solver.push()
                solver.add(Not(second == 0))
                if check_sat(solver) == unsat:
                    # it is provable that second is indeed equal to zero
                    computed = 0
                else:

                    solver.push()
                    solver.add(first < 0) # check sign of first element
                    sign = BitVecVal(-1, 256) if check_sat(solver) == sat \
                        else BitVecVal(1, 256)
                    solver.pop()

                    z3_abs = lambda x: If(x >= 0, x, -x)
                    first = z3_abs(first)
                    second = z3_abs(second)

                    computed = sign * (first % second)
                solver.pop()

            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "ADDMOD":
        if len(stack) > 2:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            third = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_third = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            taint_third_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second, third):
                if third == 0:
                    computed = 0
                else:
                    computed = (first + second) % third
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add( Not(third == 0) )
                if check_sat(solver) == unsat:
                    computed = 0
                else:
                    first = ZeroExt(256, first)
                    second = ZeroExt(256, second)
                    third = ZeroExt(256, third)
                    computed = (first + second) % third
                    computed = Extract(255, 0, computed)
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second|taint_third)
            taint_stack_prng.insert(0,taint_first_prng|taint_second_prng|taint_third_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MULMOD":
        if len(stack) > 2:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            third = stack.pop(0)

            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_third = taint_stack.pop(0)

            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            taint_third_prng = taint_stack_prng.pop(0)


            if isAllReal(first, second, third):
                if third == 0:
                    computed = 0
                else:
                    computed = (first * second) % third
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add( Not(third == 0) )
                if check_sat(solver) == unsat:
                    computed = 0
                else:
                    first = ZeroExt(256, first)
                    second = ZeroExt(256, second)
                    third = ZeroExt(256, third)
                    computed = URem(first * second, third)
                    computed = Extract(255, 0, computed)
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second|taint_third)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng | taint_third_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "EXP":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            base = stack.pop(0)
            exponent = stack.pop(0)
            taint_base = taint_stack.pop(0)
            taint_exponent = taint_stack.pop(0)
            taint_base_prng = taint_stack_prng.pop(0)
            taint_exponent_prng = taint_stack_prng.pop(0)
            # Type conversion is needed when they are mismatched
            if isAllReal(base, exponent):
                computed = pow(base, exponent, 2**256)

            elif isReal(base) and isSymbolic(exponent):
                base = BitVecVal(base, 256)
                computed = base**exponent
            elif isReal(exponent) and isSymbolic(base):
                base = BitVecVal(exponent, 256)
                computed = base**exponent
            else:
                # The computed value is unknown, this is because power is
                # not supported in bit-vector theory
                new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
                computed = BitVec(new_var_name, 256)

            check_revert = False
            if jump_type[block] == 'conditional':
                jump_target = vertices[block].get_jump_target()
                falls_to = vertices[block].get_falls_to()
                check_revert = any([True for instruction in vertices[jump_target].get_instructions() if instruction.startswith('REVERT')])
                if not check_revert:
                    check_revert = any([True for instruction in vertices[falls_to].get_instructions() if instruction.startswith('REVERT')])

            if jump_type[block] != 'conditional' or not check_revert:
                if not isAllReal(computed, base):
                    solver.push()
                    solver.add(UGT(base, computed))
                    if check_sat(solver) == sat:
                        global_problematic_pcs['integer_overflow'].append(Overflow(global_state['pc'] - 1, solver.model()))
                        overflow_pcs.append(global_state['pc'] - 1)
                    solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_base | taint_exponent)
            taint_stack_prng.insert(0, taint_base_prng|taint_exponent_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SIGNEXTEND":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if first >= 32 or first < 0:
                    computed = second
                else:
                    signbit_index_from_right = 8 * first + 7
                    if second & (1 << signbit_index_from_right):
                        computed = second | (2 ** 256 - (1 << signbit_index_from_right))
                    else:
                        computed = second & ((1 << signbit_index_from_right) - 1 )
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add( Not( Or(first >= 32, first < 0 ) ) )
                if check_sat(solver) == unsat:
                    computed = second
                else:
                    signbit_index_from_right = 8 * first + 7
                    solver.push()
                    solver.add(second & (1 << signbit_index_from_right) == 0)
                    if check_sat(solver) == unsat:
                        computed = second | (2 ** 256 - (1 << signbit_index_from_right))
                    else:
                        computed = second & ((1 << signbit_index_from_right) - 1)
                    solver.pop()
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    #
    #  10s: Comparison and Bitwise Logic Operations
    #
    elif opcode == "LT":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                first = to_unsigned(first)
                second = to_unsigned(second)
                if first < second:
                    computed = 1
                else:
                    computed = 0
            else:
                computed = If(ULT(first, second), BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng|taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "GT":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                first = to_unsigned(first)
                second = to_unsigned(second)
                if first > second:
                    computed = 1
                else:
                    computed = 0
            else:
                computed = If(UGT(first, second), BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SLT":  # Not fully faithful to signed comparison
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                first = to_signed(first)
                second = to_signed(second)
                #xiugai:taint_first = taint_stack.pop(0)
                #xiugai:taint_second = taint_stack.pop(0)
                if first < second:
                    computed = 1
                else:
                    computed = 0
            else:
                computed = If(first < second, BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SGT":  # Not fully faithful to signed comparison
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                first = to_signed(first)
                second = to_signed(second)
                if first > second:
                    computed = 1
                else:
                    computed = 0
            else:
                computed = If(first > second, BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "EQ":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if first == second:
                    computed = 1
                else:
                    computed = 0
            else:
                computed = If(first == second, BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first|taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "ISZERO":
        # Tricky: this instruction works on both boolean and integer,
        # when we have a symbolic expression, type error might occur
        # Currently handled by try and catch
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_first_prng=taint_stack_prng.pop(0)
            if isReal(first):
                if first == 0:
                    computed = 1
                else:
                    computed = 0
            elif is_bool(first):
                computed = Not(first)
            else:
                computed = If(first == 0, BitVecVal(1, 256), BitVecVal(0, 256))
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0,taint_first)
            taint_stack_prng.insert(0,taint_first_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "AND": #no need to consider symbolic value
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            computed = first & second
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "OR":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            computed = first | second
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "XOR":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            computed = first ^ second
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "NOT":
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            computed = (~first) & UNSIGNED_BOUND_NUMBER
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first)
            taint_stack_prng.insert(0, taint_first_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "BYTE":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            first = stack.pop(0)
            byte_index = 32 - first - 1
            second = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(first, second):
                if first >= 32 or first < 0:
                    computed = 0
                else:
                    computed = second & (255 << (8 * byte_index))
                    computed = computed >> (8 * byte_index)
            else:
                first = to_symbolic(first)
                second = to_symbolic(second)
                solver.push()
                solver.add( Not (Or( first >= 32, first < 0 ) ) )
                if check_sat(solver) == unsat:
                    computed = 0
                else:
                    computed = second & (255 << (8 * byte_index))
                    computed = computed >> (8 * byte_index)
                solver.pop()
            computed = simplify(computed) if is_expr(computed) else computed
            stack.insert(0, computed)
            taint_stack.insert(0, taint_first | taint_second)
            taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    #
    # 20s: SHA3
    #
    elif opcode == "SHA3":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            s0 = stack.pop(0)
            s1 = stack.pop(0)
            taint_first = taint_stack.pop(0)
            taint_second = taint_stack.pop(0)
            taint_first_prng = taint_stack_prng.pop(0)
            taint_second_prng = taint_stack_prng.pop(0)
            if isAllReal(s0, s1):
                # simulate the hashing of sha3
                data = [str(x) for x in memory[s0: s0 + s1]]
                position = ''.join(data)
                position = re.sub('[\s+]', '', position)
                position = zlib.compress(six.b(position), 9)
                position = base64.b64encode(position)
                position = position.decode('utf-8', 'strict')
                if s0 in taint_mem_prng.keys():
                    taint_stack_prng.insert(0,taint_mem_prng[s0])
                else:
                    taint_stack_prng.insert(0,SAFE_FLAG)
                #print(taint_stack_prng)
                if position in sha3_list:
                    stack.insert(0, sha3_list[position])
                    taint_stack.insert(0, taint_first | taint_second)
                    #taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
                else:
                    new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
                    new_var = BitVec(new_var_name, 256)
                    sha3_list[position] = new_var
                    stack.insert(0, new_var)
                    taint_stack.insert(0, taint_first | taint_second)
                    #taint_stack_prng.insert(0,taint_first_prng|taint_second_prng)
            else:
                # push into the execution a fresh symbolic variable
                new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var
                stack.insert(0, new_var)
                taint_stack.insert(0, taint_first | taint_second)
                taint_stack_prng.insert(0, taint_first_prng | taint_second_prng)
        else:
            raise ValueError('STACK underflow')
    #
    # 30s: Environment Information
    #
    elif opcode == "ADDRESS":  # get address of currently executing account
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, path_conditions_and_vars["Ia"])
        taint_stack.insert(0, SAFE_FLAG)
        taint_stack_prng.insert(0, SAFE_FLAG)
    elif opcode == "BALANCE":
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            address = stack.pop(0)
            taint_address = taint_stack.pop(0)
            taint_address_prng = taint_stack_prng.pop(0)

            new_var_name = gen.gen_balance_var(global_state["pc"]-1)
            if new_var_name in path_conditions_and_vars:
                new_var = path_conditions_and_vars[new_var_name]
            else:
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var
            if isReal(address):
                hashed_address = "concrete_address_" + str(address)
            else:
                hashed_address = str(address)
            global_state["balance"][hashed_address] = new_var
            stack.insert(0, new_var)
            taint_stack.insert(0,taint_address)
            taint_stack_prng.insert(0,taint_address_prng)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "CALLER":  # get caller address
        # that is directly responsible for this execution
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["sender_address"])
        taint_stack.insert(0,TAINT_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "ORIGIN":  # get execution origination address
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["origin"])
        taint_stack.insert(0, SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "CALLVALUE":  # get value of this transaction
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["value"])
        taint_stack.insert(0, SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "CALLDATALOAD":  # from input data from environment
        global_params.PRE_OPCODE_COUNT = 3
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            position = stack.pop(0)
            taint_position = taint_stack.pop(0)
            taint_position_prng = taint_stack_prng.pop(0)
            if g_src_map:
                source_code = g_src_map.get_source_code(global_state['pc'] - 1)
                if source_code.startswith("function") and isReal(position) and current_func_name in g_src_map.func_name_to_params:
                    params =  g_src_map.func_name_to_params[current_func_name]
                    param_idx = (position - 4) // 32
                    for param in params:
                        if param_idx == param['position']:
                            new_var_name = param['name']
                            g_src_map.var_names.append(new_var_name)
                else:
                    new_var_name = gen.gen_data_var(position,global_state["pc"]-1)
            else:
                new_var_name = gen.gen_data_var(position,global_state["pc"]-1)
            if new_var_name in path_conditions_and_vars:
                new_var = path_conditions_and_vars[new_var_name]
            else:
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var
            stack.insert(0, new_var)
            taint_stack.insert(0, TAINT_FLAG)
            taint_stack_prng.insert(0, SAFE_FLAG)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "CALLDATASIZE":
        global_params.PRE_OPCODE_COUNT = 3
        global_state["pc"] = global_state["pc"] + 1
        new_var_name = gen.gen_data_size(global_state["pc"] - 1)
        if new_var_name in path_conditions_and_vars:
            new_var = path_conditions_and_vars[new_var_name]
        else:
            new_var = BitVec(new_var_name, 256)
            path_conditions_and_vars[new_var_name] = new_var
        stack.insert(0, new_var)
        taint_stack.insert(0, SAFE_FLAG)
        taint_stack_prng.insert(0, SAFE_FLAG)
    elif opcode == "CALLDATACOPY":  # Copy input data to memory
        #  TODO: Don't know how to simulate this yet
        if len(stack) > 2:
            global_state["pc"] = global_state["pc"] + 1
            stack.pop(0)
            stack.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "CODESIZE":
        if g_disasm_file.endswith('.disasm'):
            evm_file_name = g_disasm_file[:-7]
        else:
            evm_file_name = g_disasm_file
        with open(evm_file_name, 'r') as evm_file:
            evm = evm_file.read()[:-1]
            code_size = len(evm)/2
            stack.insert(0, code_size)
            taint_stack.insert(0, TAINT_FLAG)
            taint_stack_prng.insert(0, SAFE_FLAG)
    elif opcode == "CODECOPY":
        if len(stack) > 2:
            global_state["pc"] = global_state["pc"] + 1
            mem_location = stack.pop(0)
            code_from = stack.pop(0)
            no_bytes = stack.pop(0)
            taint_mem_location = taint_stack.pop(0)
            taint_code_form = taint_stack.pop(0)
            taint_no_bytes = taint_stack.pop(0)
            taint_mem_location_prng = taint_stack_prng.pop(0)
            taint_code_form_prng = taint_stack_prng.pop(0)
            taint_no_bytes_prng = taint_stack_prng.pop(0)
            current_miu_i = global_state["miu_i"]

            if isAllReal(mem_location, current_miu_i, code_from, no_bytes):
                if six.PY2:
                    temp    = long(math.ceil((mem_location + no_bytes) / float(32)))
                else:
                    temp = int(math.ceil((mem_location + no_bytes) / float(32)))

                if temp > current_miu_i:
                    current_miu_i = temp

                if g_disasm_file.endswith('.disasm'):
                    evm_file_name = g_disasm_file[:-7]
                else:
                    evm_file_name = g_disasm_file
                with open(evm_file_name, 'r') as evm_file:
                    evm = evm_file.read()[:-1]
                    start = code_from * 2
                    end = start + no_bytes * 2
                    code = evm[start: end]
                mem[mem_location] = int(code, 16)
                taint_mem[mem_location] = SAFE_FLAG
                taint_mem_prng[mem_location] = SAFE_FLAG
            else:
                new_var_name = gen.gen_code_var("Ia", code_from, no_bytes, global_state["pc"]-1)
                if new_var_name in path_conditions_and_vars:
                    new_var = path_conditions_and_vars[new_var_name]
                else:
                    new_var = BitVec(new_var_name, 256)
                    path_conditions_and_vars[new_var_name] = new_var

                temp = ((mem_location + no_bytes) / 32) + 1
                current_miu_i = to_symbolic(current_miu_i)
                expression = current_miu_i < temp
                solver.push()
                solver.add(expression)
                if MSIZE:
                    if check_sat(solver) != unsat:
                        current_miu_i = If(expression, temp, current_miu_i)
                solver.pop()
                mem.clear() # very conservative
                taint_mem.clear()
                taint_mem_prng.clear()
                mem[str(mem_location)] = new_var
                taint_mem[str(mem_location)] = SAFE_FLAG
                taint_mem_prng[str(mem_location)] = SAFE_FLAG
            global_state["miu_i"] = current_miu_i
        else:
            raise ValueError('STACK underflow')
    elif opcode == "RETURNDATACOPY":
        if len(stack) > 2:
            global_state["pc"] += 1
            stack.pop(0)
            stack.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "RETURNDATASIZE":
        global_state["pc"] += 1
        new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
        new_var = BitVec(new_var_name, 256)
        stack.insert(0, new_var)
        taint_stack.insert(0,TAINT_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "GASPRICE":
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["gas_price"])
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)

    elif opcode == "EXTCODESIZE":
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            address = stack.pop(0)
            taint_address = taint_stack.pop(0)
            taint_address_prng = taint_stack_prng.pop(0)

            new_var_name = gen.gen_code_size_var(address,global_state["pc"]-1)
            if new_var_name in path_conditions_and_vars:
                new_var = path_conditions_and_vars[new_var_name]
            else:
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var
            stack.insert(0, new_var)
            taint_stack.insert(0, SAFE_FLAG)
            taint_stack_prng.insert(0, SAFE_FLAG)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "EXTCODECOPY":
        if len(stack) > 3:
            global_state["pc"] = global_state["pc"] + 1
            address = stack.pop(0)
            mem_location = stack.pop(0)
            code_from = stack.pop(0)
            no_bytes = stack.pop(0)
            current_miu_i = global_state["miu_i"]

            taint_address = taint_stack.pop(0)
            taint_mem_location = taint_stack.pop(0)
            taint_code_from = taint_stack.pop(0)
            taint_no_bytes = taint_stack.pop(0)

            taint_address_prng = taint_stack_prng.pop(0)
            taint_mem_location_prng = taint_stack_prng.pop(0)
            taint_code_form_prng = taint_stack_prng.pop(0)
            taint_no_bytes_prng = taint_stack_prng.pop(0)



            new_var_name = gen.gen_code_var(address, code_from, no_bytes,global_state["pc"]-1)
            if new_var_name in path_conditions_and_vars:
                new_var = path_conditions_and_vars[new_var_name]
            else:
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var

            temp = ((mem_location + no_bytes) / 32) + 1
            current_miu_i = to_symbolic(current_miu_i)
            expression = current_miu_i < temp
            solver.push()
            solver.add(expression)
            if MSIZE:
                if check_sat(solver) != unsat:
                    current_miu_i = If(expression, temp, current_miu_i)
            solver.pop()
            mem.clear() # very conservative
            mem[str(mem_location)] = new_var
            taint_mem.clear() # very conservative
            taint_mem_prng.clear()
            taint_mem[str(mem_location)] = SAFE_FLAG
            taint_mem_prng[str(mem_location)] = SAFE_FLAG
            global_state["miu_i"] = current_miu_i
        else:
            raise ValueError('STACK underflow')
    #
    #  40s: Block Information
    #
    elif opcode == "BLOCKHASH":  # information from block header
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            blockNumber = stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            new_var_name = "IH_blockhash"
            if new_var_name in path_conditions_and_vars:
                new_var = path_conditions_and_vars[new_var_name]
            else:
                new_var = BitVec(new_var_name, 256)
                path_conditions_and_vars[new_var_name] = new_var
            stack.insert(0, new_var)
            taint_stack.insert(0,SAFE_FLAG)
            expression1 = simplify(ULT(global_state["currentNumber"],blockNumber))
            expression2 = simplify(ULT(blockNumber,global_state["currentNumber"]-256))
            expression3 = UGT(blockNumber+256,blockNumber)
            expression4 = UGT(global_state["currentNumber"],256)
            solver.push()
            solver.add(expression3)
            solver.add(expression4)
            solver.add(simplify(Or(expression1,expression2)))
            if check_sat(solver) != unsat:
                taint_stack_prng.insert(0,TAINT_FLAG)
                solver.pop()
            else:
                solver.pop()
                if isReal(blockNumber):
                    taint_stack_prng.insert(0,TAINT_FLAG)
                else:
                    if "IH_i" in str(blockNumber):
                        mysolver = Solver()
                        mysolver.push()
                        mysolver.add(UGT(blockNumber, global_state["currentNumber"]))
                        mysolver.add(UGT(blockNumber+256,blockNumber))
                        if check_sat(mysolver) != unsat:
                            taint_stack_prng.insert(0,SAFE_FLAG)
                        else:
                            taint_stack_prng.insert(0,TAINT_FLAG)
                    elif ("Ia_store" in str(blockNumber)):
                        mysolver = Solver()
                        for var in get_vars(blockNumber):
                            varName = str(var)
                            if g_src_map and len(varName.split('-')) > 1:
                                varName = varName.split('-')[1]
                            if (varName in global_params.STORAGE_VALUES.keys()):
                                mysolver.add(var == global_params.STORAGE_VALUES[varName])
                        # for var in get_vars(blockNumber):
                        #     if(str(var) in global_params.STORAGE_VALUES.keys()):
                        #         mysolver.add(var == global_params.STORAGE_VALUES[str(var)])
                        mysolver.add(UGT(blockNumber, global_state["currentNumber"]))
                        mysolver.add(UGT(blockNumber + 256, blockNumber))
                        if check_sat(mysolver) != unsat:
                            taint_stack_prng.insert(0, SAFE_FLAG)
                        else:
                            taint_stack_prng.insert(0,TAINT_FLAG)
                    else:
                        taint_stack_prng.insert(0,TAINT_FLAG)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "COINBASE":  # information from block header
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["currentCoinbase"])
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,TAINT_FLAG)
    elif opcode == "TIMESTAMP":  # information from block header
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["currentTimestamp"])
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,TAINT_FLAG)
    elif opcode == "NUMBER":  # information from block header
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["currentNumber"])
        taint_stack.insert(0,SAFE_FLAG)
        current_block = block
        deep = 0
        while(deep < 3):
            deep = deep + 1
            if jump_type[current_block] != "terminal" and jump_type[current_block] != "unconditional":
                current_block = vertices[current_block].get_falls_to()
                if current_block:
                    check_revert = any([True for instruction in vertices[current_block].get_instructions() if
                                        (instruction.startswith('REVERT') or instruction.startswith('INVALID'))])
                    if check_revert:
                        taint_stack_prng.insert(0,SAFE_FLAG)
                        return
        taint_stack_prng.insert(0,TAINT_FLAG)
    elif opcode == "DIFFICULTY":  # information from block header
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["currentDifficulty"])
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,TAINT_FLAG)
    elif opcode == "GASLIMIT":  # information from block header
        global_state["pc"] = global_state["pc"] + 1
        stack.insert(0, global_state["currentGasLimit"])
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,TAINT_FLAG)
    #
    #  50s: Stack, Memory, Storage, and Flow Information
    #
    elif opcode == "POP":
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MLOAD":
        if len(stack) > 0:
            global_state["pc"] = global_state["pc"] + 1
            address = stack.pop(0)
            taint_address = taint_stack.pop(0)
            taint_address_prng = taint_stack_prng.pop(0)
            current_miu_i = global_state["miu_i"]
            if isAllReal(address, current_miu_i) and address in mem:
                if six.PY2:
                    temp = long(math.ceil((address + 32) / float(32)))
                else:
                    temp = int(math.ceil((address + 32) / float(32)))
                if temp > current_miu_i:
                    current_miu_i = temp
                value = mem[address]
                stack.insert(0, value)
                taint_value = taint_mem[address]
                taint_value_prng = taint_mem_prng[address]
                taint_stack.insert(0, taint_value)
                taint_stack_prng.insert(0, taint_value_prng)
            else:
                temp = ((address + 31) / 32) + 1
                current_miu_i = to_symbolic(current_miu_i)
                expression = current_miu_i < temp
                solver.push()
                solver.add(expression)
                if MSIZE:
                    if check_sat(solver) != unsat:
                        # this means that it is possibly that current_miu_i < temp
                        current_miu_i = If(expression,temp,current_miu_i)
                solver.pop()
                new_var_name = gen.gen_mem_var(address,global_state["pc"]-1)
                if new_var_name in path_conditions_and_vars:
                    new_var = path_conditions_and_vars[new_var_name]
                else:
                    new_var = BitVec(new_var_name, 256)
                    path_conditions_and_vars[new_var_name] = new_var
                stack.insert(0, new_var)
                taint_stack.insert(0, taint_address)
                taint_stack_prng.insert(0,taint_address_prng)
                if isReal(address):
                    mem[address] = new_var
                    taint_mem[address] = taint_address
                    taint_mem_prng[address] = taint_address_prng
                else:
                    mem[str(address)] = new_var
                    taint_mem[address] = taint_address
                    taint_mem_prng[address] = taint_address_prng
            global_state["miu_i"] = current_miu_i
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MSTORE":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            stored_address = stack.pop(0)
            stored_value = stack.pop(0)
            taint_stored_address = taint_stack.pop(0)
            taint_stored_value = taint_stack.pop(0)
            taint_stored_address_prng = taint_stack_prng.pop(0)
            taint_stored_value_prng = taint_stack_prng.pop(0)
            current_miu_i = global_state["miu_i"]
            if isReal(stored_address):
                # preparing data for hashing later
                old_size = len(memory) // 32
                new_size = ceil32(stored_address + 32) // 32
                mem_extend = (new_size - old_size) * 32
                memory.extend([0] * mem_extend)
                value = stored_value
                for i in range(31, -1, -1):
                    memory[stored_address + i] = value % 256
                    value /= 256
            if isAllReal(stored_address, current_miu_i):
                if six.PY2:
                    temp = long(math.ceil((stored_address + 32) / float(32)))
                else:
                    temp = int(math.ceil((stored_address + 32) / float(32)))
                if temp > current_miu_i:
                    current_miu_i = temp
                mem[stored_address] = stored_value  # note that the stored_value could be symbolic
                taint_mem[stored_address] = taint_stored_value
                taint_mem_prng[stored_address] = taint_stored_value_prng
            else:
                temp = ((stored_address + 31) / 32) + 1
                expression = current_miu_i < temp
                solver.push()
                solver.add(expression)
                if MSIZE:
                    if check_sat(solver) != unsat:
                        # this means that it is possibly that current_miu_i < temp
                        current_miu_i = If(expression,temp,current_miu_i)
                solver.pop()
                mem.clear()  # very conservative
                taint_mem.clear()
                taint_mem_prng.clear()
                mem[str(stored_address)] = stored_value
                taint_mem[str(stored_address)] = taint_stored_value
                taint_mem_prng[str(stored_address)] = taint_stored_value_prng
            global_state["miu_i"] = current_miu_i
        else:
            raise ValueError('STACK underflow')
    elif opcode == "MSTORE8":
        if len(stack) > 1:
            global_state["pc"] = global_state["pc"] + 1
            stored_address = stack.pop(0)
            temp_value = stack.pop(0)
            taint_stored_address = taint_stack.pop(0)
            taint_temp_value = taint_stack.pop(0)
            taint_stored_address_prng = taint_stack_prng.pop(0)
            taint_temp_value_prng = taint_stack_prng.pop(0)
            stored_value = temp_value % 256  # get the least byte
            current_miu_i = global_state["miu_i"]
            if isAllReal(stored_address, current_miu_i):
                if six.PY2:
                    temp = long(math.ceil((stored_address + 1) / float(32)))
                else:
                    temp = int(math.ceil((stored_address + 1) / float(32)))
                if temp > current_miu_i:
                    current_miu_i = temp
                mem[stored_address] = stored_value  # note that the stored_value could be symbolic
                taint_mem[stored_address] = taint_temp_value
                taint_mem_prng[stored_address] = taint_temp_value_prng
            else:
                temp = (stored_address / 32) + 1
                if isReal(current_miu_i):
                    current_miu_i = BitVecVal(current_miu_i, 256)
                expression = current_miu_i < temp
                solver.push()
                solver.add(expression)
                if MSIZE:
                    if check_sat(solver) != unsat:
                        # this means that it is possibly that current_miu_i < temp
                        current_miu_i = If(expression,temp,current_miu_i)
                solver.pop()
                mem.clear()  # very conservative
                taint_mem.clear()
                taint_mem_prng.clear()
                mem[str(stored_address)] = stored_value
                taint_mem[str(stored_address)] = taint_temp_value
                taint_mem_prng[str(stored_address)] = taint_temp_value_prng
            global_state["miu_i"] = current_miu_i
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SLOAD":

        if len(stack) > 0:
            path_condition_sload = path_conditions_and_vars["path_condition"]
            global_state["pc"] = global_state["pc"] + 1
            position = stack.pop(0)
            taint_position = taint_stack.pop(0)
            taint_position_prng = taint_stack_prng.pop(0)
            theKey = str(position)
            if not theKey in global_state["lastStore"]:
                if not theKey in global_state["flowSload"]:
                    global_state["flowSload"][theKey] = {}
                if not global_state['pc'] in global_state["flowSload"][theKey]:
                    global_state["flowSload"][theKey][global_state['pc']] = path_condition_sload

                global_params.BUG_TEST_CASE["TOD"][global_state["pc"]] = current_path_model
                global_params.BUG_PC["TOD"][global_state["pc"]] = str(global_state["pc"]) + ": " + str(opcode)
                global_params.BUG_BLOCK_PATH["TOD"][global_state["pc"]] = current_block_path

            if isReal(position) and position in global_state["Ia"]:
                # ADD TOD
                value = global_state["Ia"][position]
                stack.insert(0, value)
                taint_stack.insert(0,taint_position)
                taint_stack_prng.insert(0,taint_position_prng)
            else:
                if str(position) in global_state["Ia"]:
                    value = global_state["Ia"][str(position)]
                    stack.insert(0, value)
                    taint_stack.insert(0, taint_position)
                    taint_stack_prng.insert(0, taint_position_prng)
                    # theKey = str(position)
                    # if not theKey in global_state["lastStore"]:
                    #     if not theKey in global_params.GLOBAL_PC_SSLOAD:
                    #         global_params.GLOBAL_PC_SSLOAD[theKey] = {}
                    #     if not global_state['pc'] in global_params.GLOBAL_PC_SSLOAD[theKey]:
                    #         global_params.GLOBAL_PC_SSLOAD[theKey][global_state['pc']] = path_condition_sload
                    #
                    #     global_params.BUG_TEST_CASE["TOD"][global_state["pc"]] = current_path_model
                    #     global_params.BUG_PC["TOD"][global_state["pc"]] = str(global_state["pc"]) + ": " + str(opcode)
                    #     global_params.BUG_BLOCK_PATH["TOD"][global_state["pc"]] = current_block_path
                else:
                    if is_expr(position):
                        position = simplify(position)
                    if g_src_map:
                        new_var_name = g_src_map.get_source_code(global_state['pc'] - 1)
                        operators = '[-+*/%|&^!><=]'
                        new_var_name = re.compile(operators).split(new_var_name)[0].strip()
                        new_var_name = g_src_map.get_parameter_or_state_var(new_var_name)
                        if new_var_name:
                            new_var_name = gen.gen_owner_store_var(position, new_var_name,global_state["pc"]-1)
                        else:
                            new_var_name = gen.gen_owner_store_var(position,global_state["pc"]-1)
                    else:
                        new_var_name = gen.gen_owner_store_var(position,global_state["pc"]-1)

                    if new_var_name in path_conditions_and_vars:
                        new_var = path_conditions_and_vars[new_var_name]
                    else:
                        new_var = BitVec(new_var_name, 256)
                        path_conditions_and_vars[new_var_name] = new_var
                    stack.insert(0, new_var)
                    taint_stack.insert(0,taint_position)
                    taint_stack_prng.insert(0,taint_position_prng)
                    if isReal(position):
                        global_state["Ia"][position] = new_var
                    else:
                        global_state["Ia"][str(position)] = new_var
        else:
            raise ValueError('STACK underflow')

    elif opcode == "SSTORE":
        path_condition_sstore = path_conditions_and_vars["path_condition"]
        if len(stack) > 1:
            for call_pc in calls:
                calls_affect_state[call_pc] = True
            global_state["pc"] = global_state["pc"] + 1

            stored_address = stack.pop(0)

            taint_stored_address = taint_stack.pop(0)
            taint_stored_address_prng = taint_stack_prng.pop(0)
            stored_value = stack.pop(0)
            taint_stored_value = taint_stack.pop(0)

            taint_stored_value_prng = taint_stack_prng.pop(0)
            #log.info(stored_address)
            # if not stored_address in global_params.STORAGE_VALUES:
            #     global_params.STORAGE_VALUES[stored_address] = []
            # global_params.STORAGE_VALUES[stored_address].insert(0, stored_value)
            # The solvtion of key
            #xiugai:key_value = str(stack[0])
            # find the key_value instead of key address
            key_value=str(stored_value)
            if key_value.find("Ia_store") >= 0:
                ms_key = key_value
                ms_owner_key = key_value.split('-')
                try:
                    ms_owner_num = int(ms_owner_key[1])
                except:
                    ms_owner_num = str(ms_owner_key[1])
                if not(ms_owner_num in global_params.TREE):
                    global_params.TREE[ms_owner_num] = []

            if isReal(stored_address):
                # note that the stored_value could be unknown
                try:
                    stored_address = int(stored_address)
                except:
                    stored_address = str(stored_address)

                global_state["Ia"][stored_address] = stored_value
                global_params.STORAGE_VALUES[str(stored_address)] = stored_value
                if not(stored_address in global_params.TREE):
                    global_params.TREE[stored_address] = []

                flag = False
                for condition in path_condition_sstore:
                    if (str(condition).find('Is) ==') >= 0) or (str(condition).find("Extract(159, 0, Is)") >= 0):
                        flag = True
                        pc_key = str(condition)
                        break
                if flag:
                    ms_store = str(pc_key).find("Ia_store")
                    if ms_store >= 0:
                        ms_store_key = pc_key.split('-')
                        try:
                            ms_store_num = int(ms_store_key[1])
                        except:
                            ms_store_num = str(ms_store_key[1])
                        if not (ms_store_num in global_params.TREE):
                            global_params.TREE[ms_store_num] = []
                        # if ms_store_num in global_params.TREE:
                        #     if not (ms_store_num in global_params.MODIFIER):
                        #         global_params.MODIFIER.append(ms_store_num)
                        # else:
                        #     global_params.TREE[ms_store_num] = []
                        #     global_params.MODIFIER.append(ms_store_num)
                if taint_stored_value == 1:
                    # if stored_address in var_state :
                    if not flag:
                        if not (stored_address in global_params.TAINT):
                            global_params.TAINT.append(stored_address)
                    if flag and ms_store >= 0:
                        global_params.TREE[stored_address].append(ms_store_num)
                        if not (stored_address in global_params.MODIFIER):
                            global_params.MODIFIER[stored_address] = []
                            global_params.MODIFIER[stored_address].append(ms_store_num)
                        else:
                            global_params.MODIFIER[stored_address].append(ms_store_num)

                elif key_value.find("Ia_store") >= 0:
                    if flag and ms_store >= 0:
                        global_params.TREE[stored_address].append(ms_store_num)
                        if not (stored_address in global_params.MODIFIER):
                            global_params.MODIFIER[stored_address] = []
                            global_params.MODIFIER[stored_address].append(ms_store_num)
                        else:
                            global_params.MODIFIER[stored_address].append(ms_store_num)

                    global_params.TREE[stored_address].append(ms_owner_num)
            else:
                try:
                    stored_address = int(stored_address)
                except:
                    stored_address = str(stored_address)
                global_params.STORAGE_VALUES[str(stored_address)] = stored_value
                global_state["Ia"][stored_address] = stored_value
                if not (stored_address in global_params.TREE):
                    global_params.TREE[stored_address] = []

                flag = False
                for condition in path_condition_sstore:
                    if (str(condition).find('Is) ==') >= 0) or (str(condition).find("== Extract(159, 0, Is)") >= 0):
                        flag = True
                        pc_key = str(condition)
                        break
                if flag:
                    ms_store = pc_key.find("Ia_store")
                    if ms_store >= 0:
                        ms_store_key = pc_key.split('-')
                        try:
                            ms_store_num = int(ms_store_key[1])
                        except:
                            ms_store_num = str(ms_store_key[1])
                        if not (ms_store_num in global_params.TREE):
                            global_params.TREE[ms_store_num] = []
                        # if ms_store_num in global_params.TREE:
                        #     if not (ms_store_num in global_params.MODIFIER):
                        #         global_params.MODIFIER.append(ms_store_num)
                        # else:
                        #     global_params.TREE[ms_store_num] = []
                        #     global_params.MODIFIER.append(ms_store_num)
                if taint_stored_value == 1:
                    # if stored_address in var_state :
                    if not flag:
                        if not (stored_address in global_params.TAINT):
                            global_params.TAINT.append(stored_address)
                    if flag and ms_store >= 0:
                        global_params.TREE[stored_address].append(ms_store_num)
                        if not (stored_address in global_params.MODIFIER):
                            global_params.MODIFIER[stored_address] = []
                            global_params.MODIFIER[stored_address].append(ms_store_num)
                        else:
                            global_params.MODIFIER[stored_address].append(ms_store_num)

                elif key_value.find("Ia_store") >= 0:
                    if flag and ms_store >= 0:
                        global_params.TREE[stored_address].append(ms_store_num)
                        if not (stored_address in global_params.MODIFIER):
                            global_params.MODIFIER[stored_address] = []
                            global_params.MODIFIER[stored_address].append(ms_store_num)
                        else:
                            global_params.MODIFIER[stored_address].append(ms_store_num)

                    global_params.TREE[stored_address].append(ms_owner_num)
            # add tod
            theKey = str(stored_address)
            if not theKey in global_state["lastStore"]:
                global_state["lastStore"][theKey] = {}
            if not str("pc_num") in global_state["lastStore"][theKey]:
                global_state["lastStore"][theKey]["pc_num"] = 0
            if not str("path_condition") in global_state["lastStore"][theKey]:
                global_state["lastStore"][theKey]["path_condition"] = []
            global_state["lastStore"][theKey]["pc_num"] = global_state['pc']
            global_state["lastStore"][theKey]["path_condition"] = path_condition_sstore
            global_params.BUG_TEST_CASE["TOD"][global_state["pc"]] = current_path_model
            global_params.BUG_PC["TOD"][global_state["pc"]] = str(global_state["pc"]) + ": " + str(opcode)
            global_params.BUG_BLOCK_PATH["TOD"][global_state["pc"]] = current_block_path
        else:
            raise ValueError('STACK underflow')
    elif opcode == "JUMP":
        if len(stack) > 0:
            target_address = stack.pop(0)
            taint_target_address = taint_stack.pop(0)
            taint_target_address_prng = taint_stack_prng.pop(0)
            if isSymbolic(target_address):
                try:
                    target_address = int(str(simplify(target_address)))
                except:
                    raise TypeError("Target address must be an integer")
            vertices[block].set_jump_target(target_address)
            if target_address not in edges[block]:
                edges[block].append(target_address)
                vertices[target_address].set_jump_from(block)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "JUMPI":
        # We need to prepare two branches
        if len(stack) > 1:
            target_address = stack.pop(0)
            taint_target_address = taint_stack.pop(0)
            taint_target_address_prng = taint_stack_prng.pop(0)
            if isSymbolic(target_address):
                try:
                    target_address = int(str(simplify(target_address)))
                except:
                    raise TypeError("Target address must be an integer")
            vertices[block].set_jump_target(target_address)
            flag = stack.pop(0)
            taint_flag = taint_stack.pop(0)
            taint_flag_prng = taint_stack_prng.pop(0)
            branch_expression = (BitVecVal(0, 1) == BitVecVal(1, 1))
            if isReal(flag):
                if flag != 0:
                    branch_expression = True
            elif is_bool(flag):
                branch_expression = flag
            else:
                branch_expression = (flag != 0)
            vertices[block].set_branch_expression(branch_expression)
            vertices[block].set_taint_branch_expression(taint_flag_prng)
            if target_address not in edges[block]:
                edges[block].append(target_address)
                vertices[target_address].set_jump_from(block)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "PC":
        stack.insert(0, global_state["pc"])
        global_state["pc"] = global_state["pc"] + 1
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "MSIZE":
        global_state["pc"] = global_state["pc"] + 1
        msize = 32 * global_state["miu_i"]
        stack.insert(0, msize)
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "GAS":
        # In general, we do not have this precisely. It depends on both
        # the initial gas and the amount has been depleted
        # we need o think about this in the future, in case precise gas
        # can be tracked
        global_state["pc"] = global_state["pc"] + 1
        new_var_name = gen.gen_gas_var(global_state["pc"]-1)
        new_var = BitVec(new_var_name, 256)
        path_conditions_and_vars[new_var_name] = new_var
        stack.insert(0, new_var)
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
    elif opcode == "JUMPDEST":
        # Literally do nothing
        global_state["pc"] = global_state["pc"] + 1
    #
    #  60s & 70s: Push Operations
    #
    elif opcode.startswith('PUSH', 0):  # this is a push instruction
        position = int(opcode[4:], 10)
        global_state["pc"] = global_state["pc"] + 1 + position
        pushed_value = int(instr_parts[1], 16)
        stack.insert(0, pushed_value)
        taint_stack.insert(0,SAFE_FLAG)
        taint_stack_prng.insert(0,SAFE_FLAG)
        if global_params.UNIT_TEST == 3: # test evm symbolic
            stack[0] = BitVecVal(stack[0], 256)
            taint_stack[0] = taint_stack[0]
            taint_stack_prng[0] = taint_stack_prng[0]

    #
    #  80s: Duplication Operations
    #
    elif opcode.startswith("DUP", 0):
        global_state["pc"] = global_state["pc"] + 1
        position = int(opcode[3:], 10) - 1
        if len(stack) > position:
            duplicate = stack[position]
            taint_duplicate = taint_stack[position]
            taint_duplicate_prng = taint_stack_prng[position]
            stack.insert(0, duplicate)
            taint_stack.insert(0,taint_duplicate)
            taint_stack_prng.insert(0,taint_duplicate_prng)
        else:
            raise ValueError('STACK underflow')

    #
    #  90s: Swap Operations
    #
    elif opcode.startswith("SWAP", 0):
        global_state["pc"] = global_state["pc"] + 1
        position = int(opcode[4:], 10)
        if len(stack) > position:
            temp = stack[position]
            stack[position] = stack[0]
            stack[0] = temp
            taint_temp = taint_stack[position]
            taint_stack[position] = taint_stack[0]
            taint_stack[0] = taint_temp
            taint_temp_prng = taint_stack_prng[position]
            taint_stack_prng[position] = taint_stack_prng[0]
            taint_stack_prng[0] = taint_temp_prng
        else:
            raise ValueError('STACK underflow')

    #
    #  a0s: Logging Operations
    #
    elif opcode in ("LOG0", "LOG1", "LOG2", "LOG3", "LOG4"):
        global_state["pc"] = global_state["pc"] + 1
        # We do not simulate these log operations
        num_of_pops = 2 + int(opcode[3:])
        while num_of_pops > 0:
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            num_of_pops -= 1

    #
    #  f0s: System Operations
    #
    elif opcode == "CREATE":
        if len(stack) > 2:
            global_state["pc"] += 1
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
            new_var = BitVec(new_var_name, 256)
            stack.insert(0, new_var)
            taint_stack.insert(0, SAFE_FLAG)
            taint_stack_prng.insert(0,SAFE_FLAG)
        else:
            raise ValueError('STACK underflow')
    elif opcode == "CALL":
        # TODO: Need to handle miu_i
        recipient_pos = -1
        amount_pos = -1
        path_condition_call = path_conditions_and_vars["path_condition"]
        if len(stack) > 6:
            calls.append(global_state["pc"])
            for call_pc in calls:
                if call_pc not in calls_affect_state:
                    calls_affect_state[call_pc] = False
            global_state["pc"] = global_state["pc"] + 1
            outgas = stack.pop(0)
            taint_outgas =  taint_stack.pop(0)
            taint_outgas_prng = taint_stack_prng.pop(0)
            recipient = stack.pop(0)
            taint_recipient = taint_stack.pop(0)
            taint_recipient_prng = taint_stack_prng.pop(0)
            transfer_amount = stack.pop(0)
            if isSymbolic(recipient) and is_storage_var(recipient):

                recipient_pos = get_storage_position(recipient)
                theKey = str(recipient_pos)
                if not theKey in global_state["lastStore"]:
                    if not theKey in global_params.GLOBAL_PC_SSLOAD:
                        global_params.GLOBAL_PC_SSLOAD[theKey] = {}
                    if not global_state['pc'] in global_params.GLOBAL_PC_SSLOAD[theKey]:
                        global_params.GLOBAL_PC_SSLOAD[theKey][global_state['pc']] = path_condition_call
            if isSymbolic(transfer_amount) and is_storage_var(transfer_amount):

                amount_pos = get_storage_position(transfer_amount)
                theKey = str(amount_pos)
                if not theKey in global_state["lastStore"]:
                    if not theKey in global_params.GLOBAL_PC_SSLOAD:
                        global_params.GLOBAL_PC_SSLOAD[theKey] = {}
                    if not global_state['pc'] in global_params.GLOBAL_PC_SSLOAD[theKey]:
                        global_params.GLOBAL_PC_SSLOAD[theKey][global_state['pc']] = path_condition_call

            for item in path_condition_call:
                if is_expr(item):
                    if str(item).find("Ia_store") >= 0:
                        pc_store = str(item)
                        pc_key = pc_store.split('-')[1]

                        if pc_key in global_state["flowSload"]:
                            if not pc_key in global_params.GLOBAL_PC_SSLOAD:
                                global_params.GLOBAL_PC_SSLOAD[pc_key] = {}

                            for item_pc in global_state["flowSload"][pc_key]:
                                if not item_pc in global_params.GLOBAL_PC_SSLOAD[pc_key]:
                                    global_params.GLOBAL_PC_SSLOAD[pc_key][item_pc] = []
                                global_params.GLOBAL_PC_SSLOAD[pc_key][item_pc] = global_state["flowSload"][pc_key][item_pc]

            taint_transfer_amount = taint_stack.pop(0)
            taint_transfer_amount_prng = taint_stack_prng.pop(0)
            start_data_input = stack.pop(0)
            taint_start_data_input = taint_stack.pop(0)
            taint_start_data_input_prng = taint_stack_prng.pop(0)
            size_data_input = stack.pop(0)
            taint_size_date_input = taint_stack.pop(0)
            taint_size_data_input_prng = taint_stack_prng.pop(0)
            start_data_output = stack.pop(0)
            taint_start_data_output = taint_stack.pop(0)
            taint_start_data_output_prng = taint_stack_prng.pop(0)
            size_data_ouput = stack.pop(0)
            taint_size_data_output = taint_stack.pop(0)
            taint_size_data_ouput_prng = taint_stack_prng.pop(0)
            # in the paper, it is shaky when the size of data output is
            # min of stack[6] and the | o |

            if isReal(transfer_amount):
                if transfer_amount == 0:
                    stack.insert(0, 1)   # x = 0
                    taint_stack.insert(0, SAFE_FLAG)
                    taint_stack_prng.insert(0, SAFE_FLAG)
                    return

            # Let us ignore the call depth
            balance_ia = global_state["balance"]["Ia"]
            is_enough_fund = (transfer_amount <= balance_ia)
            solver.push()
            solver.add(is_enough_fund)

            if check_sat(solver) == unsat:
                # this means not enough fund, thus the execution will result in exception
                solver.pop()
                stack.insert(0, 0)   # x = 0
                taint_stack.insert(0, SAFE_FLAG)
                taint_stack_prng.insert(0,SAFE_FLAG)
            else:
                # the execution is possibly okay
                if isReal(recipient):
                    taint_recipient=1
                return_call_status = gen.gen_call_return_var(taint_recipient,global_state["pc"]-1)
                returnVar = BitVec(return_call_status,256)
                path_conditions_and_vars[return_call_status]=returnVar
                stack.insert(0, returnVar)   # x = 1
                taint_stack.insert(0, SAFE_FLAG)
                taint_stack_prng.insert(0,SAFE_FLAG)
                solver.pop()
                solver.add(is_enough_fund)
                path_conditions_and_vars["path_condition"].append(is_enough_fund)
                last_idx = len(path_conditions_and_vars["path_condition"]) - 1
                new_balance_ia = (balance_ia - transfer_amount)
                global_state["balance"]["Ia"] = new_balance_ia
                address_is = path_conditions_and_vars["Is"]
                address_is = (address_is & CONSTANT_ONES_159)
                boolean_expression = (recipient != address_is)
                solver.push()
                solver.add(boolean_expression)
                if check_sat(solver) == unsat:
                    solver.pop()
                    new_balance_is = (global_state["balance"]["Is"] + transfer_amount)
                    global_state["balance"]["Is"] = new_balance_is
                else:
                    solver.pop()
                    if isReal(recipient):
                        new_address_name = "concrete_address_" +str(recipient)+"_"+str(global_state["pc"]-1)
                    else:
                        new_address_name = gen.gen_arbitrary_address_var(global_state["pc"]-1)
                    old_balance_name = gen.gen_arbitrary_var(global_state["pc"]-1)
                    old_balance = BitVec(old_balance_name, 256)
                    path_conditions_and_vars[old_balance_name] = old_balance
                    constraint = (old_balance >= 0)
                    solver.add(constraint)
                    path_conditions_and_vars["path_condition"].append(constraint)
                    new_balance = (old_balance + transfer_amount)
                    global_state["balance"][new_address_name] = new_balance
        else:
            raise ValueError('STACK underflow')
    elif opcode == "CALLCODE":
        # TODO: Need to handle miu_i
        if len(stack) > 6:
            calls.append(global_state["pc"])
            for call_pc in calls:
                if call_pc not in calls_affect_state:
                    calls_affect_state[call_pc] = False
            global_state["pc"] = global_state["pc"] + 1
            outgas = stack.pop(0)
            recipient = stack.pop(0) # this is not used as recipient
            if global_params.USE_GLOBAL_STORAGE:
                if isReal(recipient):
                    recipient = hex(recipient)
                    if recipient[-1] == "L":
                        recipient = recipient[:-1]
                    recipients.add(recipient)
                else:
                    recipients.add(None)

            transfer_amount = stack.pop(0)
            start_data_input = stack.pop(0)
            size_data_input = stack.pop(0)
            start_data_output = stack.pop(0)
            size_data_ouput = stack.pop(0)
            # in the paper, it is shaky when the size of data output is
            # min of stack[6] and the | o |
            taint_outgas = taint_stack.pop(0)
            taint_recipient = taint_stack.pop(0)  #
            taint_transfer_amount = taint_stack.pop(0)
            taint_start_data_input = taint_stack.pop(0)
            taint_size_data_input = taint_stack.pop(0)
            taint_start_data_output = taint_stack.pop(0)
            taint_size_data_ouput = taint_stack.pop(0)

            taint_outgas_prng = taint_stack_prng.pop(0)
            taint_recipient_prng = taint_stack_prng.pop(0)  #
            taint_transfer_amount_prng = taint_stack_prng.pop(0)
            taint_start_data_input_prng = taint_stack_prng.pop(0)
            taint_size_data_input_prng = taint_stack_prng.pop(0)
            taint_start_data_output_prng = taint_stack_prng.pop(0)
            taint_size_data_ouput_prng = taint_stack_prng.pop(0)


            if isReal(transfer_amount):
                if transfer_amount == 0:
                    stack.insert(0, 1)   # x = 0
                    #xiugai:
                    taint_stack.insert(0,SAFE_FLAG)
                    taint_stack_prng.insert(0,SAFE_FLAG)
                    return

            # Let us ignore the call depth
            balance_ia = global_state["balance"]["Ia"]
            is_enough_fund = (transfer_amount <= balance_ia)
            solver.push()
            solver.add(is_enough_fund)

            if check_sat(solver) == unsat:
                # this means not enough fund, thus the execution will result in exception
                solver.pop()
                stack.insert(0, 0)   # x = 0
                taint_stack.insert(0,SAFE_FLAG)
                taint_stack_prng.insert(0,SAFE_FLAG)
            else:
                # the execution is possibly okay
                stack.insert(0, 1)   # x = 1
                taint_stack.insert(0, SAFE_FLAG)
                taint_stack_prng.insert(0,SAFE_FLAG)
                solver.pop()
                solver.add(is_enough_fund)
                path_conditions_and_vars["path_condition"].append(is_enough_fund)
                last_idx = len(path_conditions_and_vars["path_condition"]) - 1
        else:
            raise ValueError('STACK underflow')
    elif opcode in ("DELEGATECALL", "STATICCALL"):
        if len(stack) > 5:
            global_state["pc"] += 1
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            recipient = stack.pop(0)
            taint_recipient = taint_stack.pop(0)
            taint_recipient_prng=taint_stack_prng.pop(0)
            if global_params.USE_GLOBAL_STORAGE:
                if isReal(recipient):
                    recipient = hex(recipient)
                    if recipient[-1] == "L":
                        recipient = recipient[:-1]
                    recipients.add(recipient)
                else:
                    recipients.add(None)

            stack.pop(0)
            stack.pop(0)
            stack.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            new_var_name = gen.gen_arbitrary_var(global_state["pc"]-1)
            new_var = BitVec(new_var_name, 256)
            stack.insert(0, new_var)
            taint_stack.insert(0,SAFE_FLAG)
            taint_stack_prng.insert(0,SAFE_FLAG)
        else:
            raise ValueError('STACK underflow')
    elif opcode in ("RETURN", "REVERT"):
        # TODO: Need to handle miu_i
        if len(stack) > 1:
            if opcode == "REVERT":
                revertible_overflow_pcs.update(overflow_pcs)
                revertible_overflow_pcs.update(underflow_pcs)
                global_state["pc"] = global_state["pc"] + 1
            stack.pop(0)
            stack.pop(0)
            taint_stack.pop(0)
            taint_stack.pop(0)
            taint_stack_prng.pop(0)
            taint_stack_prng.pop(0)
            # TODO
            pass
        else:
            raise ValueError('STACK underflow')
    elif opcode == "SUICIDE":
        global_state["pc"] = global_state["pc"] + 1
        recipient = stack.pop(0)
        taint_recipient_prng = taint_stack_prng.pop(0)
        taint_recipient = taint_stack.pop(0)
        transfer_amount = global_state["balance"]["Ia"]
        global_state["balance"]["Ia"] = 0
        if isReal(recipient):
            new_address_name = "concrete_address_" + str(recipient)+"_"+str(global_state["pc"]-1)
        else:
            new_address_name = gen.gen_arbitrary_address_var(global_state["pc"]-1)
        old_balance_name = gen.gen_arbitrary_var(global_state["pc"]-1)
        old_balance = BitVec(old_balance_name, 256)
        path_conditions_and_vars[old_balance_name] = old_balance
        constraint = (old_balance >= 0)
        solver.add(constraint)
        path_conditions_and_vars["path_condition"].append(constraint)
        new_balance = (old_balance + transfer_amount)
        global_state["balance"][new_address_name] = new_balance
        # TODO
        return

    else:
        log.debug("UNKNOWN INSTRUCTION: " + opcode)
        if global_params.UNIT_TEST == 2 or global_params.UNIT_TEST == 3:
            log.critical("Unknown instruction: %s" % opcode)
            exit(UNKNOWN_INSTRUCTION)
        raise Exception('UNKNOWN INSTRUCTION: ' + opcode)


def check_fun(condition):
    if str(condition).find("Extract(255, 224") >= 0:
        return False
    else:
        return True
# detect transaction order

def detect_tod_bug():
    global results
    global g_src_map
    global tod_bug

    tod_solver = Solver()
    tod_solver.set("timeout", global_params.TIMEOUT)
    tod_list = {}
    for sstore_adr in global_params.GLOBAL_PC_SSTORE:
        # sstore_lenth = len(global_params.GLOBAL_PC_SSTORE[sstore_adr])
        all_pc = []
        for store_keys in global_params.GLOBAL_PC_SSTORE[sstore_adr]:
            all_pc.append(store_keys)
            path_condition_store = global_params.GLOBAL_PC_SSTORE[sstore_adr][store_keys]
            if sstore_adr in global_params.GLOBAL_PC_SSLOAD:
                for sload_keys in global_params.GLOBAL_PC_SSLOAD[sstore_adr]:
                    path_condition_sload = global_params.GLOBAL_PC_SSLOAD[sstore_adr][sload_keys]
                    tod_solver.push()
                    store_function = []
                    sload_function = []
                    for single_sload_condition in path_condition_sload:
                        check = check_fun(single_sload_condition)
                        if check:
                            tod_solver.add(single_sload_condition)
                        else:
                            sload_function = str(single_sload_condition)
                    for single_store_condition in path_condition_store:

                        check = check_fun(single_store_condition)
                        if check:
                            tod_solver.add(single_store_condition)
                        else:
                            store_function = str(single_store_condition)
                    if len(store_function) != 0 and len(sload_function) != 0 :
                        if store_function == sload_function:
                            continue

                    if check_sat(tod_solver) == sat:
                        if not (store_keys in tod_list):
                            tod_list[store_keys] = []
                        tod_list[store_keys].append(sload_keys)

                    tod_solver.pop()

        # for i in range(len(all_pc) -1):
        #     for j in range(i+1, len(all_pc)):
        #
        #
        #         path_condition_i = global_params.GLOBAL_PC_SSTORE[sstore_adr][all_pc[i]]
        #         path_condition_j = global_params.GLOBAL_PC_SSTORE[sstore_adr][all_pc[j]]
        #
        #
        #         tod_solver.push()
        #         for single_i_condition in path_condition_i:
        #             check = check_fun(single_i_condition)
        #             if check:
        #                 tod_solver.add(single_i_condition)
        #             tod_solver.add(single_i_condition)
        #         for single_j_condition in path_condition_j:
        #             check = check_fun(single_j_condition)
        #             if check:
        #                 tod_solver.add(single_j_condition)
        #         if check_sat(tod_solver,False) == sat:
        #             if not (all_pc[i] in tod_list):
        #                 tod_list[all_pc[i]] = []
        #             tod_list[all_pc[i]].append(all_pc[j])
        #
        #     tod_solver.pop()
    real_tod = []

    for key in tod_list:
        real_tod.append(key)
        for value in tod_list[key]:
            real_tod.append(value)
            #global_params.BUG_PC["TOD"][key] = global_params.BUG_PC["TOD"][key] + 'dependency'+ str(value)
            #global_params.BUG_PC["TOD"][tod_list[key]] = global_params.BUG_PC["TOD"][value] + 'dependency' + str(key)
    tod_bug = todBug(g_src_map, real_tod)

    if g_src_map:
        results['vulnerabilities']['tod'] = tod_bug.get_warnings()
    else:
        results['vulnerabilities']['tod'] = tod_bug.is_vulnerable()
    log.info("\t Transaction Ordering Dependency: \t\t %s", tod_bug.is_vulnerable())
    if tod_bug.is_vulnerable():
        tod_bug.set_paths(global_params.BUG_BLOCK_PATH["TOD"])
        tod_bug.set_test_case(global_params.BUG_TEST_CASE["TOD"])
        tod_bug.set_types(global_params.BUG_PC["TOD"])
        # log.info("\t  Reentrancy Numbers: \t\t %s", len(reentrancy.get_pcs()))
        global_params.REP_BUG_NUM["Transaction-Order Dependency"] = len(real_tod)


def detect_integer_underflow():
    global integer_underflow

    underflows_pc = []
    for underflow in global_problematic_pcs['integer_underflow']:
        if underflow.pc not in revertible_underflow_pcs:
            underflows_pc.append(underflow.pc)
            global_params.BUG_PC["UNDERFLOW"][underflow.pc] = str(underflow.pc) + " : " + str(underflow.opcode) + " : Integer Underflow"
            global_params.BUG_BLOCK_PATH["UNDERFLOW"][underflow.pc] = underflow.path
            unmodelString = {}
            for variable in underflow.model.decls():
                unmodelString[str(variable)]=str(underflow.model[variable])
            global_params.BUG_TEST_CASE["UNDERFLOW"][underflow.pc] = unmodelString
    integer_underflow = IntegerUnderflow(g_src_map, underflows_pc)

    if g_src_map:
        results['vulnerabilities']['integer_underflow'] = integer_underflow.get_warnings()
    else:
        results['vulnerabilities']['integer_underflow'] = integer_underflow.is_vulnerable()

    log.info('\t  Integer Underflow: \t\t %s', integer_underflow.is_vulnerable())
    if not global_params.EVAL:
        if integer_underflow.is_vulnerable():
            integer_underflow.set_paths(global_params.BUG_BLOCK_PATH["UNDERFLOW"])
            integer_underflow.set_test_case(global_params.BUG_TEST_CASE["UNDERFLOW"])
            integer_underflow.set_types(global_params.BUG_PC["UNDERFLOW"])
            # log.info("\t  IntegerOverflow Numbers: \t\t %s", len(integer_overflow.get_pcs()))
            global_params.REP_BUG_NUM["Integer Underflow"] = len(integer_underflow.get_pcs())
            # for key in integer_overflow.get_pcs():
            #     log.info("\t  Bug abstract: \t\t %s", integer_overflow.get_types()[key])
            #     print_dot_cfg_path(integer_overflow.get_paths()[key],integer_overflow.get_pc_test_case()[key])
            #     log.info("\t  Bug test case: \t\t %s", integer_overflow.get_pc_test_case()[key])

def detect_integer_overflow():
    global integer_overflow

    overflows_pc = []
    for overflow in global_problematic_pcs['integer_overflow']:
        if overflow.pc not in revertible_overflow_pcs:
            overflows_pc.append(overflow.pc)
            global_params.BUG_PC["OVERFLOW"][overflow.pc] = str(overflow.pc) + " : " + str(overflow.opcode) + " : Integer Overflow"
            global_params.BUG_BLOCK_PATH["OVERFLOW"][overflow.pc] = overflow.path
            modelString = {}
            for variable in overflow.model.decls():
                modelString[str(variable)]=str(overflow.model[variable])
            global_params.BUG_TEST_CASE["OVERFLOW"][overflow.pc] = modelString
    integer_overflow = IntegerOverflow(g_src_map, overflows_pc)

    if g_src_map:
        results['vulnerabilities']['integer_overflow'] = integer_overflow.get_warnings()
    else:
        results['vulnerabilities']['integer_overflow'] = integer_overflow.is_vulnerable()
    log.info('\t  Integer Overflow: \t\t %s', integer_overflow.is_vulnerable())
    if not global_params.EVAL:
        if integer_overflow.is_vulnerable():
            integer_overflow.set_paths(global_params.BUG_BLOCK_PATH["OVERFLOW"])
            integer_overflow.set_test_case(global_params.BUG_TEST_CASE["OVERFLOW"])
            integer_overflow.set_types(global_params.BUG_PC["OVERFLOW"])
            # log.info("\t  IntegerOverflow Numbers: \t\t %s", len(integer_overflow.get_pcs()))
            global_params.REP_BUG_NUM["Integer Overflow/Underflow"] = len(integer_overflow.get_pcs())
            # for key in integer_overflow.get_pcs():
            #     log.info("\t  Bug abstract: \t\t %s", integer_overflow.get_types()[key])
            #     print_dot_cfg_path(integer_overflow.get_paths()[key],integer_overflow.get_pc_test_case()[key])
            #     log.info("\t  Bug test case: \t\t %s", integer_overflow.get_pc_test_case()[key])


#yzq:
def detect_prng_bug():
    global g_src_map
    global results
    global prng_bug

    prng_bug = PrngBug(g_src_map, global_problematic_pcs['prng_bug'])



    if g_src_map:
        results['vulnerabilities']['prng_bug'] = prng_bug.get_warnings()
    else:
        results['vulnerabilities']['prng_bug'] = prng_bug.is_vulnerable()
    log.info("\t  prng Vulnerability: \t\t %s", prng_bug.is_vulnerable())
    if not global_params.EVAL:
        if prng_bug.is_vulnerable():
            prng_bug.set_paths(global_params.BUG_BLOCK_PATH["PRNG"])
            prng_bug.set_test_case(global_params.BUG_TEST_CASE["PRNG"])
            prng_bug.set_types(global_params.BUG_PC["PRNG"])
            # log.info("\t  prng Numbers: \t\t %s", len(prng_bug.get_pcs()))
            global_params.REP_BUG_NUM["Pseudorandom Generator"] = len(prng_bug.get_pcs())
            # for key in prng_bug.get_pcs():
            #     log.info("\t  Bug abstract: \t\t %s", prng_bug.get_types()[key])
            #     print_dot_cfg_path(prng_bug.get_paths()[key],prng_bug.get_pc_test_case()[key])
            #     log.info("\t  Bug test case: \t\t %s", prng_bug.get_pc_test_case()[key])


def detect_dos_bug():
    global g_src_map
    global results
    global dos_bug
    global vertices
    global edges
    global jump_type

    real_dos = []
    for gpc in global_problematic_pcs['dos_bug'].keys():
        # for parent in vertices[global_problematic_pcs["dos_bug"][gpc]].get_jump_from():
        #     for brother in edges[parent]:
        #         if brother != global_problematic_pcs["dos_bug"][gpc] and jump_type[brother] == "terminal":
        #             real_dos.append(gpc)
        if "loop gas limit" in global_params.BUG_PC['DOS'][gpc]:
            real_dos.append(gpc)
            continue
        for child in edges[global_problematic_pcs["dos_bug"][gpc]]:
            check_revert = any([True for instruction in vertices[child].get_instructions() if
                                instruction.startswith('REVERT')])
            check_RETURNDATACOPY = any([True for instruction in vertices[child].get_instructions() if
                                instruction.startswith('RETURNDATACOPY')])
            if child != global_problematic_pcs["dos_bug"][gpc] and check_revert and not check_RETURNDATACOPY: #require()
                real_dos.append(gpc)
            elif child != global_problematic_pcs["dos_bug"][gpc] and check_revert and check_RETURNDATACOPY:  #tranfer()
                if gpc-3 in instructions:
                    instr = instructions[gpc-3]
                    pc = instr.split(' ')[1]
                    pc = int(pc, 16)
                    real_dos.append(pc+1)
                    global_params.BUG_BLOCK_PATH["DOS"][pc+1]=global_params.BUG_BLOCK_PATH["DOS"][gpc]
                    global_params.BUG_TEST_CASE["DOS"][pc+1]=global_params.BUG_TEST_CASE["DOS"][gpc]
                    global_params.BUG_PC["DOS"][pc+1]=global_params.BUG_PC["DOS"][gpc]
                else:
                    real_dos.append(gpc)
    dos_bug = DosBug(g_src_map, real_dos)

    if g_src_map:
        results['vulnerabilities']['dos_bug'] = dos_bug.get_warnings()
    else:
        results['vulnerabilities']['dos_bug'] = dos_bug.is_vulnerable()
    #print results
    log.info("\t  DOS Vulnerability: \t\t %s", dos_bug.is_vulnerable())
    if not global_params.EVAL:
        if dos_bug.is_vulnerable():
            dos_bug.set_paths(global_params.BUG_BLOCK_PATH["DOS"])
            dos_bug.set_test_case(global_params.BUG_TEST_CASE["DOS"])
            dos_bug.set_types(global_params.BUG_PC["DOS"])
            #log.info("\t  DOS Numbers: \t\t %s", len(dos_bug.get_pcs()))
            global_params.REP_BUG_NUM["Denial-of-Service"] = len(dos_bug.get_pcs())
            # for key in dos_bug.get_pcs():
            #     log.info("\t  Bug abstract: \t\t %s", dos_bug.get_types()[key])
            #     print_dot_cfg_path(dos_bug.get_paths()[key], dos_bug.get_pc_test_case()[key])
            #     log.info("\t  Bug test case: \t\t %s", dos_bug.get_pc_test_case()[key])

def detect_assertion_failure():
    global g_src_map
    global results
    global assertion_failure

    assertion_failure = AssertionFailure(g_src_map, global_problematic_pcs['assertion_failure'])
    if g_src_map:
        results['vulnerabilities']['assertion_failure'] = assertion_failure.get_warnings()
    else:
        results['vulnerabilities']['assertion_failure'] = assertion_failure.is_vulnerable()
    s = "\t  Account Integrity Violation: \t\t\t %s" % assertion_failure.is_vulnerable()
    log.info(s)
    if not global_params.EVAL:
        if assertion_failure.is_vulnerable():
            assertion_failure.set_paths(global_params.BUG_BLOCK_PATH["ASSERTFAIL"])
            assertion_failure.set_test_case(global_params.BUG_TEST_CASE["ASSERTFAIL"])
            assertion_failure.set_types(global_params.BUG_PC["ASSERTFAIL"])
            #log.info("\t  DOS Numbers: \t\t %s", len(dos_bug.get_pcs()))
            global_params.REP_BUG_NUM["Account Integrity Violation"] = len(assertion_failure.get_pcs())

def detect_vulnerabilities():
    global results
    global g_src_map
    global visited_pcs
    global global_problematic_pcs
    global begin
    global reentrancy

    if instructions:
        evm_code_coverage = float(len(visited_pcs)) / len(instructions.keys()) * 100
        log.info("\t  EVM Code Coverage: \t\t\t %s%%", round(evm_code_coverage, 1))
        results["evm_code_coverage"] = str(round(evm_code_coverage, 1))

        detect_integer_overflow()
        detect_integer_underflow()
        detect_prng_bug()
        detect_dos_bug()
        detect_tod_bug()
        log.info("\t Transaction Ordering Dependency: \t\t %s", False)
        detect_assertion_failure()

        real_reentrancy = []

        if len(global_params.TAINT) != 0 or len(global_params.D_TAINT) != 0:
            for item in global_params.TARGET:

                dfs_results = 0
                if len(global_params.TREE[item]) != 0:
                    dfs_results = dfs_target(item, global_params.TARGET_DEPTH, global_params.MODIFIER_DEPTH)
                    # if dfs_results == 2:
                    #     flag = True
                    # elif dfs_results == 1:
                    #     flag = True
                if item in global_params.D_TAINT:
                    dfs_results = 3
                    # log.info("\t  Reentrancy Vulnerability:\t\tTrue")
                    # log.info("taint direct happen")
                    # flag = True
                if item in global_params.TAINT:
                    dfs_results = 4
                    # log.info("\t  Reentrancy Vulnerability: \t\t True")
                    # log.info("taint in-direct happen")
                    # flag = True
                if dfs_results:
                    for single in global_params.TARGET_PC[item]:
                        #xiugai
                        if not single in global_params.BUG_PC["REENTRANCY"]:
                            continue
                        real_reentrancy.append(single)
                        if dfs_results == 2:
                            if (item in global_params.D_MODIFIER) and (len(global_params.TREE[item]) == 1) and (
                                    global_params.TREE[item][0] in global_params.D_MODIFIER[item]):
                                global_params.BUG_PC["REENTRANCY"][single] = global_params.BUG_PC["REENTRANCY"][single] + 'taint target owner direct happen and onlyowner not work'
                            else:
                                global_params.BUG_PC["REENTRANCY"][single] = global_params.BUG_PC["REENTRANCY"][single] + 'taint target owner direct happen and onlyowner not work'
                        elif dfs_results == 1:
                            global_params.BUG_PC["REENTRANCY"][single] = global_params.BUG_PC["REENTRANCY"][single] + 'target taint transfer'
                        elif dfs_results == 3:
                            global_params.BUG_PC["REENTRANCY"][single] = global_params.BUG_PC["REENTRANCY"][single] + 'taint direct happen'
                        elif dfs_results == 4:
                            global_params.BUG_PC["REENTRANCY"][single] = global_params.BUG_PC["REENTRANCY"][single] + 'taint in-direct happen'

        reentrancy = Reentrancy(g_src_map, real_reentrancy)

        if g_src_map:
            results['vulnerabilities']['reentrancy'] = reentrancy.get_warnings()
        else:
            results['vulnerabilities']['reentrancy'] = reentrancy.is_vulnerable()
        log.info("\t  Reentrancy: \t\t %s", reentrancy.is_vulnerable())
        if reentrancy.is_vulnerable():
            reentrancy.set_paths(global_params.BUG_BLOCK_PATH["REENTRANCY"])
            reentrancy.set_test_case(global_params.BUG_TEST_CASE["REENTRANCY"])
            reentrancy.set_types(global_params.BUG_PC["REENTRANCY"])
            # log.info("\t  Reentrancy Numbers: \t\t %s", len(reentrancy.get_pcs()))
            global_params.REP_BUG_NUM["Reentrancy"] = len(reentrancy.get_pcs())
            # for key in reentrancy.get_pcs():
            #     log.info("\t  Bug abstract: \t\t %s", reentrancy.get_types()[key])
            #     print_dot_cfg_path(reentrancy.get_paths()[key], reentrancy.get_pc_test_case()[key])
            #     log.info("\t  Bug test case: \t\t %s", reentrancy.get_pc_test_case()[key])
    else:
        log.info("\t  NO CONTRACT")

    return results, 0
def dfs_target(item,target_time,owner_time):
    if target_time == 0:
        return 0
    for node in global_params.TREE[item]:
        
        if (node in global_params.TAINT) and (item in global_params.MODIFIER) and (node in global_params.MODIFIER[item]):
            return 2
        elif node in global_params.TAINT:
            return 1
        elif (item in global_params.MODIFIER) and (node in global_params.MODIFIER[item]):
            result = dfs_modfier(node, owner_time)
            if result:
                return 2
        else:
            result = dfs_target(node,target_time-1,owner_time)
            if result != 0:
                return result
    return 0



def dfs_modfier(node, ownertime):
    if ownertime == 0:
        # target_time = global_params.MODIFIER_DEPTH
        return False
    for item in global_params.TREE[node]:

        if item in global_params.TAINT:
            return True
        elif (node in global_params.MODIFIER) and (item in global_params.MODIFIER[node]):
            result = dfs_modfier(item, ownertime-1)
            if result:
                return True
        else:
            result_target = dfs_target(item, ownertime-1,global_params.MODIFIER_DEPTH)
            if result_target:
                return True
    return False

def overview():
    global integer_overflow
    #global reentrancy
    global prng_bug
    global dos_bug
    #global money_concurrency
    global reentrancy
    global integer_overflow
    global tod_bug
    global assertion_failure

    document = Document()
    document.add_heading('Security Verification Report')
    document.add_heading('Overview', level = 1)

    records = (
        ('Reentrancy',global_params.REP_BUG_NUM["Reentrancy"]),
        ('Pseudorandom Generator',global_params.REP_BUG_NUM["Pseudorandom Generator"]),
        ('Integer Overflow',global_params.REP_BUG_NUM["Integer Overflow"]),
        ('Integer Underflow', global_params.REP_BUG_NUM["Integer Underflow"]),
        ('Transaction-Order Dependency',global_params.REP_BUG_NUM["Transaction-Order Dependency"]),
        ('Denial-of-Service',global_params.REP_BUG_NUM["Denial-of-Service"]),
        ('Account Integrity Violation',global_params.REP_BUG_NUM["Account Integrity Violation"])
    )
    table = document.add_table(rows= 1, cols = 2)
    document.add_paragraph('')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerability Type'
    hdr_cells[1].text = '#Num'

    for vulner, num in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(vulner)
        row_cells[1].text = str(num)

    document.add_heading('Details', level=1)
    if global_params.REP_BUG_NUM["Reentrancy"]:
        case_detail(reentrancy,document,1)
    if global_params.REP_BUG_NUM["Pseudorandom Generator"]:
        case_detail(prng_bug,document,2)
    if global_params.REP_BUG_NUM["Integer Overflow"]:
        case_detail(integer_overflow,document,3)
    # if global_params.REP_BUG_NUM["Transaction-Order Dependency"]:
    #     case_detail(tod_bug,document,4)
    if global_params.REP_BUG_NUM["Denial-of-Service"]:
        case_detail(dos_bug,document,5)
    if global_params.REP_BUG_NUM["Integer Underflow"]:
        case_detail(integer_underflow,document,7)
    if global_params.REP_BUG_NUM["Account Integrity Violation"]:
        case_detail(assertion_failure,document,6)
    global_params.DOC_PATH = os.getcwd()+ separator+"reports"
    if not os.path.exists(global_params.DOC_PATH):
        os.mkdir(global_params.DOC_PATH)
    filename = global_params.reportName
    if filename == "":
        time1 = datetime.datetime.now()
        filename = datetime.datetime.strftime(time1, '%Y-%m-%d %H:%M:%S')
    document.save(global_params.DOC_PATH + separator + filename + ".docx")
    return filename

def signReport(filename):
    with open(global_params.DOC_PATH + separator+filename+'.docx', 'r') as f:
        message = f.read()
    h = MD5.new(message)
    # print(global_params.PRIVATE_KEY)
    private_key = RSA.importKey(global_params.PRIVATE_KEY)
    signer = PKCS1_v1_5.new(private_key)
    signature = signer.sign(h)
    with open(global_params.DOC_PATH + separator+filename+'signature.ca', 'w') as f:
        f.write(signature)

def case_detail(bug_type,document,bug_id):
    for key in bug_type.get_pcs():
        if bug_id == 1:
            document.add_paragraph('Case: Reentrancy', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['REENTRANCY'][key],global_params.BUG_TEST_CASE['REENTRANCY'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 2:
            document.add_paragraph('Case: Pseudorandom Generator', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['PRNG'][key],global_params.BUG_TEST_CASE['PRNG'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 3:
            document.add_paragraph('Case: Integer Overflow', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['OVERFLOW'][key],global_params.BUG_TEST_CASE['OVERFLOW'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 4:
            document.add_paragraph('Case: Transaction-Ordering Dependency', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['TOD'][key],global_params.BUG_TEST_CASE['TOD'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 5:
            document.add_paragraph('Case: Denial-of-Service', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['DOS'][key],global_params.BUG_TEST_CASE['DOS'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 6:
            document.add_paragraph('Case: Account Integrity Violation', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['ASSERTFAIL'][key],global_params.BUG_TEST_CASE['ASSERTFAIL'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))
        elif bug_id == 7:
            document.add_paragraph('Case: Integer Underflow', style='List Number')
            document.add_paragraph('Bug Abstract: ' + str(bug_type.get_types()[key]))
            path=print_dot_cfg_path(global_params.BUG_BLOCK_PATH['UNDERFLOW'][key],global_params.BUG_TEST_CASE['UNDERFLOW'][key])
            document.add_picture(path, width=Inches(3.0))
            document.add_paragraph('Bug test case: ' + str(bug_type.report_test_case(key)))

def vulnerability_found():
    global g_src_map
    global time_dependency
    global callstack
    global money_concurrency
    global assertion_failure
    global parity_multisig_bug_2

    vulnerabilities = [callstack, money_concurrency, time_dependency]

    if g_src_map and global_params.CHECK_ASSERTIONS:
        vulnerabilities.append(assertion_failure)
        vulnerabilities.append(parity_multisig_bug_2)

    for vul in vulnerabilities:
        if vul.is_vulnerable():
            return 1
    return 0

def closing_message():
    global g_disasm_file
    global results

    log.info("\t====== Analysis Completed ======")
    if global_params.EVAL:
        result_file = g_disasm_file.split('.evm.disasm')[0] + '.json'
        with open(result_file, 'w') as of:
            of.write(json.dumps(results, indent=1))
        log.info("Wrote results to %s.", result_file)

class TimeoutError(Exception):
    pass

class Timeout:
   """Timeout class using ALARM signal."""

   def __init__(self, sec=10, error_message=os.strerror(errno.ETIME)):
       self.sec = sec
       self.error_message = error_message

   def __enter__(self):
       signal.signal(signal.SIGALRM, self._handle_timeout)
       signal.alarm(self.sec)

   def __exit__(self, *args):
       signal.alarm(0)    # disable alarm

   def _handle_timeout(self, signum, frame):
       raise TimeoutError(self.error_message)

def do_nothing():
    pass

def run_build_cfg_and_analyze(timeout_cb=do_nothing):
    initGlobalVars()
    global g_timeout

    try:
        with Timeout(sec=global_params.GLOBAL_TIMEOUT):
            build_cfg_and_analyze()
        log.debug('Done Symbolic execution')
    except TimeoutError:
        g_timeout = True
        timeout_cb()

def get_recipients(disasm_file, contract_address):
    global recipients
    global g_src_map
    global g_disasm_file
    global g_source_file

    g_src_map = None
    g_disasm_file = disasm_file
    g_source_file = None
    recipients = set()

    evm_code_coverage = float(len(visited_pcs)) / len(instructions.keys())

    run_build_cfg_and_analyze()

    return {
        'addrs': list(recipients),
        'evm_code_coverage': evm_code_coverage,
        'timeout': g_timeout
    }

def test():
    global_params.GLOBAL_TIMEOUT = global_params.GLOBAL_TIMEOUT_TEST

    def timeout_cb():
        traceback.print_exc()
        exit(EXCEPTION)

    run_build_cfg_and_analyze(timeout_cb=timeout_cb)

def analyze():
    def timeout_cb():
        if global_params.DEBUG_MODE:
            traceback.print_exc()

    run_build_cfg_and_analyze(timeout_cb=timeout_cb)

def run(disasm_file=None, source_file=None, source_map=None):
    global g_disasm_file
    global g_source_file
    global g_src_map
    global results

    g_disasm_file = disasm_file
    g_source_file = source_file
    g_src_map = source_map

    if is_testing_evm():
        return test()
    else:
        begin = time.time()
        log.info("\t============ Results ===========")
        analyze()
        ret,exit_code = detect_vulnerabilities()
        if not global_params.EVAL:
            filename = overview()
            signReport(filename)
        closing_message()
        # print_visited_edges()
        global_params.TARGET = []
        global_params.MODIFIER = {}
        global_params.TAINT = []
        global_params.TREE ={}
        global_params.global_state = {}
        global_params.TARGET_PC = {}
        return ret,exit_code
